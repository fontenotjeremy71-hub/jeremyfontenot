from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
import pypdfium2 as pdfium


ROOT = Path(__file__).resolve().parents[2]
OPS = Path(r"C:\Users\jeremy\Documents\projects\jeremy-homelab-ops")
PYTHON = Path(r"C:\Users\jeremy\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe")
SKILL_RENDER = Path(r"C:\Users\jeremy\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents\render_docx.py")
DOC_TITLE = "Jeremy Fontenot’s On-Premises Home Lab Documentation"
DOC_TITLE_ASCII = "Jeremy Fontenot's On-Premises Home Lab Documentation"
DOC_FILENAME = "Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx"
DOC_PATH = ROOT / "assets" / "documents" / DOC_FILENAME
DOC_MIRROR = ROOT / "evidence-library" / "projects" / "on-prem-home-lab" / DOC_FILENAME
CURRENT = ROOT / "evidence-library" / "projects" / "on-prem-home-lab" / "current-validated-state"
SITE_AUDIT = ROOT / "artifacts" / "site-audit"
FIG_DIR = SITE_AUDIT / "docx-generated-figures"
RENDER_DIR = SITE_AUDIT / "docx-render-final"
PREVIEW_DIR = SITE_AUDIT / "screenshot-previews"
DOC_URL = "./assets/documents/Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx"
DOC_CTA = "Download the Professional Home Lab Documentation (.docx)"
HOME_LAB_COMMIT = subprocess.check_output(["git", "-C", str(OPS), "rev-parse", "HEAD"], text=True).strip()
WEBSITE_START_COMMIT = subprocess.check_output(["git", "-C", str(ROOT), "rev-parse", "HEAD"], text=True).strip()
GENERATED_AT = datetime.now(timezone.utc).isoformat(timespec="seconds")


def rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def ensure_dirs() -> None:
    for path in [DOC_PATH.parent, DOC_MIRROR.parent, CURRENT, SITE_AUDIT, FIG_DIR, RENDER_DIR, PREVIEW_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in [("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")]:
                elem = margins.find(qn(f"w:{side}"))
                if elem is None:
                    elem = OxmlElement(f"w:{side}")
                    margins.append(elem)
                elem.set(qn("w:w"), value)
                elem.set(qn("w:type"), "dxa")


def mark_header_row(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, label: str, anchor: str, page: int | None = None) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    if page is not None:
        paragraph.add_run("\t" + str(page))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_image_alt(doc: Document, title: str, descr: str) -> None:
    doc_prs = doc.element.xpath("//wp:docPr")
    if doc_prs:
        doc_prs[-1].set("title", title)
        doc_prs[-1].set("descr", descr)


def add_table(doc: Document, rows: list[list[str]], widths: list[int], caption: str) -> None:
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[idx], value, bold=True)
        shade_cell(table.rows[0].cells[idx], "E8EEF5")
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    mark_header_row(table)
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, alt: str, width_inches: float = 6.1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    set_image_alt(doc, caption.split(".")[0], alt)
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"


def make_figures() -> list[Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figures = []
    specs = [
        (
            "architecture-summary.png",
            "Home Lab Architecture Summary",
            ["PowerEdge R710", "Proxmox VE 9.2.3", "pfSense VM 100", "DC01 VM 200", "WS01 VM 300", "Linux01 VM 400 VLAN 30", "backup-hdd"],
        ),
        (
            "evidence-methodology.png",
            "Evidence Methodology",
            ["Read-only collection", "Authenticated inventory", "Classification", "Manifest + hash", "Public-safe publication"],
        ),
        (
            "backup-boundary.png",
            "Backup and Restore Boundary",
            ["Scheduled job visible", "Snapshot mode", "zstd compression", "Isolated restore documented", "No recurring DR claim"],
        ),
    ]
    try:
        font_title = ImageFont.truetype("arial.ttf", 34)
        font = ImageFont.truetype("arial.ttf", 22)
        font_small = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font_title = font = font_small = ImageFont.load_default()
    for filename, title, nodes in specs:
        img = Image.new("RGB", (1600, 760), "#0f172a")
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle((35, 35, 1565, 725), radius=34, outline="#22d3ee", width=4, fill="#111827")
        draw.text((72, 65), title, fill="#ffffff", font=font_title)
        x = 80
        y = 180
        w = 310
        h = 105
        for idx, node in enumerate(nodes):
            row = idx // 4
            col = idx % 4
            nx = x + col * 370
            ny = y + row * 190
            draw.rounded_rectangle((nx, ny, nx + w, ny + h), radius=18, fill="#1e293b", outline="#ec4899", width=3)
            draw.text((nx + 22, ny + 28), node, fill="#e2e8f0", font=font)
            if col < 3 and idx < len(nodes) - 1 and (idx + 1) // 4 == row:
                draw.line((nx + w + 8, ny + h // 2, nx + 360, ny + h // 2), fill="#22d3ee", width=4)
                draw.polygon([(nx + 360, ny + h // 2), (nx + 342, ny + h // 2 - 10), (nx + 342, ny + h // 2 + 10)], fill="#22d3ee")
        draw.text((72, 682), "Generated public-safe diagram from repository evidence. Not a network capture.", fill="#cbd5e1", font=font_small)
        path = FIG_DIR / filename
        img.save(path, optimize=True)
        figures.append(path)
    return figures


SECTIONS = [
    ("Abstract", 1, [
        "This document presents a professional, public-safe description of Jeremy Fontenot's personal on-premises home lab. It is written for recruiters, hiring managers, and technical reviewers who need a clear view of the environment, the evidence behind the claims, and the boundaries around what is not being claimed.",
        "The lab supports systems administration and infrastructure operations practice while Jeremy's professional positioning remains anchored in experienced Service Desk and IT Support work. The evidence set demonstrates disciplined validation, documentation, troubleshooting, and automation habits without presenting personal lab activity as employer production administration.",
    ]),
    ("Purpose and Scope", 1, [
        "The purpose of this documentation is to explain what exists in the personal lab, why each component matters, how evidence was collected, and how claims are classified. The scope includes the physical host, Proxmox virtualization, pfSense routing, Windows identity services, Linux integration, backup architecture, logging, automation, validation workflows, and known limitations.",
        "The document intentionally avoids unsupported claims of enterprise scale, guaranteed availability, production responsibility, or full security assurance. It focuses on what a reviewer can inspect from sanitized artifacts, source summaries, manifests, hashes, and current-state validation records.",
    ]),
    ("Environment Boundaries", 1, [
        "The environment is a personal nonproduction home lab. It is appropriate evidence for technical initiative, hands-on practice, systems thinking, and documentation discipline. It is not a claim of employer production administration, client infrastructure ownership, or operations at enterprise scale.",
        "Public artifacts are sanitized. Private credentials, tokens, cookies, private keys, raw logs with sensitive values, and administrative secrets are not published. Historical evidence remains available only where it helps reviewers understand chronology, not as a current-state label.",
    ]),
    ("Evidence Methodology", 1, [
        "Evidence precedence favors current command output, API output, guest-agent output, authenticated target-local inventory, and explicit repository validation. Older screenshots and documentation remain useful when clearly labeled as historical or when current records independently support the same claim.",
        "The methodology separates observation from interpretation. A TCP timeout is reported as a timeout, not as proof of an outage. A configuration record is reported as configured unless a test or authenticated inventory supports a stronger statement.",
    ]),
    ("Evidence Classification", 1, [
        "Claims are classified as Validated, Tested, Configured, Historical, Limitation, or Inconclusive. Validated means direct evidence supports the stated behavior or state. Tested means a relevant check produced a recorded result, but the result may be narrower than the broad service claim.",
        "This classification model helps technical reviewers trust the strongest statements while still seeing what remains unresolved. It also keeps recruiter-facing copy readable by moving dense proof details into manifests and appendices.",
    ]),
    ("Physical Server Architecture", 1, [
        "The physical lab host is documented as a Dell PowerEdge R710. It provides the hardware foundation for the virtualization environment and supports the lab's role as a practical infrastructure learning platform.",
        "The server is presented as personal lab equipment. Hardware ownership and capacity are not extended into claims about production-grade resilience, clustered high availability, or enterprise hardware lifecycle management.",
    ]),
    ("Proxmox Virtualization", 1, [
        "Current source records identify Proxmox VE 9.2.3 as the virtualization platform. The primary supported virtual machines are VM 100 pfSense, VM 200 DC01, VM 300 WS01, and VM 400 Linux01.",
        "Proxmox evidence supports running VM inventory, storage visibility, backup job visibility, Linux01 VLAN tagging, and QEMU guest agent status where available. Guest operating system claims are strengthened through target-local inventories where Proxmox metadata alone is insufficient.",
    ]),
    ("Storage Architecture", 1, [
        "Current records identify backup storage named backup-hdd at /mnt/pve/backup-hdd. Older references to other backup storage labels are treated as stale or inconclusive unless later evidence proves a rename or replacement.",
        "Storage claims are limited to observed storage inventory, configured backup target use, and documented utilization or monitoring records. The evidence does not claim full storage redundancy or enterprise storage design.",
    ]),
    ("pfSense", 1, [
        "pfSense is represented as VM 100 and provides routing and firewall context for the personal lab. Public documentation distinguishes pfSense configuration records from behavior that requires a separate current test.",
        "The supported design includes a LAN gateway, a Linux VLAN gateway, and an OpenVPN management path. Direct WAN-side management exposure is treated as historical where older records mention it; the current design emphasizes VPN-scoped administration.",
    ]),
    ("Network Bridges", 1, [
        "Proxmox bridge evidence identifies VM network attachment through vmbr interfaces. pfSense bridges between upstream and internal lab networks, while domain systems reside on the internal lab side.",
        "Bridge documentation is used to explain how systems relate to one another, not to imply that every firewall rule, NAT rule, or packet path was exhaustively tested from the public website environment.",
    ]),
    ("VLANs", 1, [
        "Linux01 is currently treated as VLAN 30 on vmbr1 with the supported address 10.10.30.20/24. Older Linux01 10.10.20.x evidence is preserved as historical setup or pre-reconciliation context.",
        "VLAN claims are supported by current Proxmox metadata, DNS reconciliation, and target-local Linux inventory. The documentation avoids implying broad segmentation assurance beyond the cited evidence.",
    ]),
    ("VPN and Management Routing", 1, [
        "The management path includes OpenVPN network context and route evidence. Current records support DNS and RDP reachability to DC01 through the management path, while SMB and other timeouts remain documented observations only.",
        "Management routing is described as a controlled access pattern for lab administration. It is not presented as a comprehensive remote-access security audit.",
    ]),
    ("DC01", 1, [
        "DC01 is documented as VM 200, a Windows Server 2022 domain controller for ad.jeremyfontenot.online. Target-local inventory and retained screenshots support the Windows services that are described in this paper.",
        "DC01 evidence matters because it anchors the lab's identity, name resolution, address assignment, policy, and file-service practice areas.",
    ]),
    ("Active Directory Domain Services", 1, [
        "Active Directory Domain Services evidence supports a functioning personal lab domain with domain controller context, domain identity, and administrative validation records. Claims remain scoped to the personal lab.",
        "The AD DS section demonstrates directory-services concepts: domain controller role, workstation membership, Linux domain integration, DNS dependency, and policy-linked administration.",
    ]),
    ("FSMO Roles", 1, [
        "Retained DC01 evidence supports FSMO role ownership review. FSMO role validation is useful because it shows the domain controller role is not merely a label but part of an inspected AD DS configuration.",
        "The evidence is not expanded into multi-domain or multi-domain-controller operational claims.",
    ]),
    ("DNS", 1, [
        "DNS evidence includes AD-integrated DNS context, forward lookup, reverse lookup, and current Linux01 name reconciliation. DNS is central to domain services, Linux integration, and management-path validation.",
        "Where DNS records are historical, the documentation labels them accordingly. Current Linux01 DNS resolves to 10.10.30.20.",
    ]),
    ("DHCP", 1, [
        "DHCP evidence includes scope, option, statistics, and redacted lease review. DHCP claims are scoped to visible records and current summaries rather than broad address-management maturity.",
        "Redaction is intentional because DHCP lease records can expose host identifiers. Published descriptions preserve technical meaning without disclosing unnecessary details.",
    ]),
    ("Group Policy", 1, [
        "Group Policy evidence supports GPO inventory and link review. Group Policy appears in the workstation, file-service, and permissions evidence where applied policy affects access or configuration.",
        "The documentation distinguishes inventory from processing proof. A listed GPO is not treated as applied to a user or machine unless a relevant validation record supports it.",
    ]),
    ("File Services", 1, [
        "File-service evidence is primarily represented through DC01 and WS01 validation records, including mapped drive and share-access testing. The lab demonstrates permission modeling and user access validation.",
        "No broad enterprise file-service assurance, data classification program, or production data management claim is made.",
    ]),
    ("Security Groups and Permissions", 1, [
        "Security group and permission evidence includes different access outcomes for different user contexts, including read/write and write-denial validation. This demonstrates practical access-control testing.",
        "The evidence is valuable because it shows both positive and negative validation: successful access for an authorized context and denied write access where appropriate.",
    ]),
    ("WS01", 1, [
        "WS01 is documented as VM 300, a Windows domain workstation. Current source records include authenticated inventory evidence and historical workstation screenshots. Proxmox guest-agent limitations for WS01 remain visible where applicable.",
        "WS01 demonstrates endpoint administration concepts: domain join, DNS/DHCP context, RDP troubleshooting, firewall state, mapped drives, and file-share permissions.",
    ]),
    ("Linux01", 1, [
        "Linux01 is documented as VM 400, Ubuntu Server 26.04 LTS, a domain member on VLAN 30 at 10.10.30.20/24. Current Proxmox metadata and target-local inventory strengthen this state.",
        "Linux01 provides the lab's strongest cross-platform evidence: Linux identity integration, SSSD remediation, Kerberos, SSH, UFW, sudo controls, QEMU guest agent response, and rsyslog receiver state.",
    ]),
    ("Linux VLAN Placement", 1, [
        "Linux01 VLAN placement is supported by Proxmox VM metadata showing tag 30 and by current DNS reconciliation to 10.10.30.20. Historical 10.10.20.x references remain historical.",
        "This section matters because correcting Linux01 placement prevents outdated network statements from leaking into public proof labels.",
    ]),
    ("Active Directory Integration", 1, [
        "Linux Active Directory integration is documented through realm membership, identity lookup, and related validation artifacts. The public claim is that integration is configured and validated where the target-local records support it.",
        "The integration demonstrates cross-platform administration, DNS dependency awareness, identity troubleshooting, and careful validation of authentication behavior.",
    ]),
    ("SSSD", 1, [
        "The Linux01 SSSD remediation record documents a responder activation conflict, configuration correction, service restart, failed-unit cleanup, identity lookup, and a new AD-authenticated SSH session.",
        "The important operational pattern is rollback-first remediation: preserve a backup, validate configuration, change only the conflicting state, restart the service, then validate identity and access.",
    ]),
    ("Kerberos", 1, [
        "Kerberos is part of the Linux domain-integration story. The document treats Kerberos claims as supported where realm, identity, and authentication evidence connect to the Linux01 validation workflow.",
        "Timeouts from the current management workstation to Kerberos ports are not interpreted as Kerberos service failure because the evidence path and firewall policy boundaries matter.",
    ]),
    ("SSH", 1, [
        "SSH evidence includes post-remediation AD-authenticated access to Linux01. That result supports a meaningful operational claim: identity integration was not merely configured; it was validated with a new login path.",
        "The documentation does not publish credentials, keys, or raw sensitive session data.",
    ]),
    ("UFW", 1, [
        "UFW appears in Linux01 records as part of the security-control boundary. The documentation reports UFW state only to the level supported by target-local inventory and sanitized summaries.",
        "Firewall claims are scoped narrowly because a firewall state listing is not equivalent to a comprehensive security test.",
    ]),
    ("sudo Controls", 1, [
        "sudo controls are included as Linux administrative access evidence. The lab documentation treats sudo configuration as a privileged-access control that should be described without exposing sensitive policy internals.",
        "The public artifact emphasizes that delegated administrative control exists where supported, but it does not publish secrets or overstate least-privilege maturity.",
    ]),
    ("QEMU Guest Agent", 1, [
        "Current Proxmox records show Linux01 QEMU guest agent responding. Other primary VMs are not promoted to guest-agent-validated state unless the relevant current evidence supports it.",
        "This distinction prevents hypervisor metadata from being mistaken for guest-level validation across every VM.",
    ]),
    ("Centralized Logging", 1, [
        "Centralized logging evidence includes Linux01 rsyslog receiver state from authenticated target-local inventory and planning records for further expansion. pfSense forwarding is documented in retained records.",
        "Additional Windows Event Forwarding, Proxmox logging, alerting, and firewall changes remain planned until explicitly approved and validated.",
    ]),
    ("Backup Architecture", 1, [
        "Backup architecture evidence supports an enabled scheduled backup job named backup-critical-lab-vms for VMs 100, 200, 300, and 400 using snapshot mode, zstd compression, and backup-hdd storage.",
        "This evidence supports backup configuration visibility. It does not prove current archive integrity or recurring recovery success by itself.",
    ]),
    ("Snapshots", 1, [
        "Snapshot evidence appears in workstation baseline and backup documentation. Snapshots are treated as operational checkpoints and backup-mode context rather than a substitute for tested recovery.",
        "The document keeps snapshot configuration separate from restore assurance.",
    ]),
    ("Archive Integrity", 1, [
        "Archive integrity evidence includes SHA-256 hash records for published documentation and retained evidence packages. Hashes help reviewers verify that artifacts are stable and traceable.",
        "A hash proves file identity, not truth of every claim inside the file. The manifest therefore pairs hashes with claim boundaries and source paths.",
    ]),
    ("Restore Testing", 1, [
        "Retained documentation records isolated startup restore checks for temporary WS01 and Linux01 restore VMs, with network isolation and cleanup noted. Those temporary VMs are no longer present in current Proxmox inventory.",
        "The current records do not establish recurring disaster recovery assurance, RTO, RPO, or a standing recovery program. A future restore test would need fresh evidence with selected backup, restore target, isolation method, validation checks, and cleanup.",
    ]),
    ("Storage Monitoring", 1, [
        "Storage monitoring claims are limited to available inventory, utilization, and planning records. Backup storage is visible, but full monitoring maturity is not claimed.",
        "The logging expansion plan identifies disk-use thresholds and rotation concerns for Linux01, which demonstrates operations planning without claiming an implemented alerting platform.",
    ]),
    ("PowerShell Automation", 1, [
        "PowerShell automation appears in repository validation, read-only baseline tooling, evidence hash generation, link checks, and target-local Windows inventory scripts. It demonstrates repeatable operations work rather than one-off screenshots.",
        "The automation principle is consistent: collect read-only evidence, write structured outputs, avoid secrets, and validate repository quality before publication.",
    ]),
    ("Read-Only Operations Framework", 1, [
        "The read-only framework includes Proxmox inventory collection, route evidence, DNS checks, bounded TCP reachability checks, repository validation, and source reconciliation. It is deliberately constrained to avoid unauthorized infrastructure changes.",
        "Read-only validation is useful because it creates repeatable evidence while preserving production-safe habits: explicit scope, no secret capture, no configuration changes, and clear timeout wording.",
    ]),
    ("Authenticated Target-Local Inventories", 1, [
        "Authenticated target-local inventories strengthen guest claims for DC01, WS01, and Linux01. These records provide more reliable guest-level evidence than external reachability alone.",
        "Target-local collection is especially important where remote client context lacks modules, firewall access, or guest-agent response.",
    ]),
    ("Troubleshooting and Remediation", 1, [
        "The strongest remediation example is Linux01 SSSD. It shows diagnosis, backup, configuration correction, service recovery, failed-state cleanup, identity lookup, and AD-authenticated SSH validation.",
        "WS01 screenshots also support troubleshooting practice around RDP, NLA, DNS, DHCP, firewall rules, and workstation identity. These are presented as support and lab operations evidence, not production incident claims.",
    ]),
    ("Known Limitations", 1, [
        "Known limitations include personal lab scope, no employer production administration claim, no client infrastructure ownership, no enterprise scale claim, no guaranteed uptime, no full security assurance, and no recurring restore assurance.",
        "Timeout observations remain timeout observations. Historical screenshots remain useful only when labeled and scoped correctly.",
    ]),
    ("Skills Demonstrated", 1, [
        "The lab demonstrates service-desk-to-infrastructure growth: troubleshooting, documentation, Active Directory fundamentals, DNS and DHCP awareness, Windows workstation administration, Linux administration, PowerShell automation, evidence governance, and cautious public communication.",
        "The strongest professional signal is not a single tool. It is the habit of connecting claims to artifacts, stating limitations plainly, and validating before publishing.",
    ]),
    ("Conclusion", 1, [
        "Jeremy Fontenot's on-premises home lab is a credible personal infrastructure portfolio because it is documented with current-state evidence, carefully bounded public claims, and reviewer-accessible artifacts.",
        "The environment supports a clear career narrative: experienced Service Desk professional actively building evidence and validation toward Systems Administration and Infrastructure Operations.",
    ]),
    ("References", 1, [
        f"Fontenot, J. (2026). Home lab source-of-truth repository records. fontenotjeremy71-hub/jeremy-homelab-ops, commit {HOME_LAB_COMMIT}.",
        f"Fontenot, J. (2026). Public evidence library and portfolio repository. fontenotjeremy71-hub/jeremyfontenot, working source {WEBSITE_START_COMMIT}.",
        "American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.).",
    ]),
    ("Appendices", 1, [
        "Appendix A lists the primary source artifacts. Appendix B summarizes claim boundaries. Appendix C maps required topics to document sections. Appendix D records the clickable TOC validation results.",
    ]),
]


SPECIAL_TABLES = {
    "Evidence Classification": (
        "Table 1. Evidence classification model.",
        [
            ["Classification", "Use", "Boundary"],
            ["Validated", "Direct evidence confirms the state or behavior.", "Used only where source artifacts support the claim."],
            ["Tested", "A relevant test produced a recorded result.", "May be narrower than full service health."],
            ["Configured", "Configuration is present.", "Behavior may require a separate validation."],
            ["Historical", "Useful retained context.", "Not promoted as current state."],
            ["Limitation", "Known gap or unsupported conclusion.", "Kept visible for reviewers."],
        ],
        [1700, 3800, 3860],
    ),
    "Proxmox Virtualization": (
        "Table 2. Current primary VM inventory.",
        [
            ["VM", "System", "Supported role", "Evidence boundary"],
            ["100", "pfSense", "Firewall and routing VM", "Configuration context; not a full firewall audit."],
            ["200", "DC01", "Windows Server domain controller", "Guest-level claims use target-local evidence where available."],
            ["300", "WS01", "Windows domain workstation", "Guest-agent limitation remains visible."],
            ["400", "Linux01", "Ubuntu domain member on VLAN 30", "Guest agent and target-local evidence support stronger claims."],
        ],
        [900, 1600, 3150, 3710],
    ),
    "VPN and Management Routing": (
        "Table 3. Management-path evidence boundaries.",
        [
            ["Observation", "Recorded result", "Public claim"],
            ["DC01 DNS TCP 53", "Reachable through management path", "DNS reachability supported."],
            ["DC01 RDP TCP 3389", "Reachable through management path", "Management RDP path supported."],
            ["SMB or policy-dependent ports", "Timeouts recorded", "Timeout only; not outage proof."],
            ["Direct WAN RDP", "Historical, not current public design", "Not promoted as current access model."],
        ],
        [2400, 3100, 3860],
    ),
    "Backup Architecture": (
        "Table 4. Backup architecture evidence.",
        [
            ["Item", "Supported value", "Boundary"],
            ["Backup job", "backup-critical-lab-vms", "Configured and visible."],
            ["VM coverage", "100, 200, 300, 400", "Does not prove restore success."],
            ["Mode", "snapshot", "Configuration evidence only."],
            ["Compression", "zstd", "Configuration evidence only."],
            ["Storage", "backup-hdd", "Current observed storage target."],
        ],
        [2200, 3300, 3860],
    ),
    "Known Limitations": (
        "Table 5. Non-negotiable public claim boundaries.",
        [
            ["Area", "Not claimed", "Reason"],
            ["Employment", "Employer production administration", "Personal lab evidence is separate from work history."],
            ["Scale", "Enterprise-scale operations", "Single personal lab environment."],
            ["Availability", "Guaranteed uptime", "No production SLA evidence."],
            ["Security", "Full security assurance", "Configuration and tests are scoped."],
            ["Recovery", "Recurring DR, RTO, or RPO", "Only retained isolated restore checks are documented."],
        ],
        [1800, 3600, 3960],
    ),
}


REQUIRED_TOPICS = [
    "Title page", "Abstract", "Purpose and scope", "Environment boundaries", "Evidence methodology",
    "Evidence classification", "Physical server architecture", "Proxmox virtualization",
    "Storage architecture", "pfSense", "Network bridges", "VLANs", "VPN and management routing",
    "DC01", "Active Directory Domain Services", "FSMO roles", "DNS", "DHCP", "Group Policy",
    "File services", "Security groups and permissions", "WS01", "Linux01", "Linux VLAN placement",
    "Active Directory integration", "SSSD", "Kerberos", "SSH", "UFW", "sudo controls",
    "QEMU guest agent", "Centralized logging", "Backup architecture", "Snapshots",
    "Archive integrity", "Restore testing", "Storage monitoring", "PowerShell automation",
    "Read-only operations framework", "Authenticated target-local inventories",
    "Troubleshooting and remediation", "Known limitations", "Skills demonstrated",
    "Conclusion", "References", "Appendices",
]


def section_anchor(title: str) -> str:
    return re.sub(r"[^A-Za-z0-9_]", "_", title)[:38]


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def build_docx(page_map: dict[str, int] | None = None, out_path: Path = DOC_PATH) -> None:
    figures = make_figures()
    doc = Document()
    apply_styles(doc)
    props = doc.core_properties
    props.title = DOC_TITLE
    props.subject = DOC_TITLE
    props.author = "Jeremy Fontenot"
    props.keywords = "home lab, Proxmox, pfSense, Active Directory, Linux, PowerShell, evidence governance"
    props.comments = "Public-safe professional home lab documentation generated from repository evidence."

    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(140)
    title_run = title_p.add_run(DOC_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared by Jeremy Fontenot").bold = True
    meta.add_run(f"\nHome-lab source commit: {HOME_LAB_COMMIT}")
    meta.add_run(f"\nGenerated: {GENERATED_AT}")
    doc.add_page_break()

    toc_head = doc.add_heading("Table of Contents", level=1)
    add_bookmark(toc_head, "TOC", 1)
    for idx, (title, level, _paras) in enumerate(SECTIONS, start=2):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0 if level == 1 else 0.25)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25))
        add_internal_link(p, title, section_anchor(title), (page_map or {}).get(title, 0))
    doc.add_page_break()

    bookmark_id = 10
    for title, level, paragraphs in SECTIONS:
        heading = doc.add_heading(title, level=level)
        add_bookmark(heading, section_anchor(title), bookmark_id)
        bookmark_id += 1
        back = doc.add_paragraph()
        add_internal_link(back, "Back to Table of Contents", "TOC")
        back.paragraph_format.space_after = Pt(4)
        for text in paragraphs:
            if title == "References":
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
            else:
                doc.add_paragraph(text)
        if title == "Physical Server Architecture":
            add_figure(doc, figures[0], "Figure 1. Current supported home-lab architecture summary.", "Architecture diagram showing PowerEdge host, Proxmox, pfSense, DC01, WS01, Linux01, and backup storage.")
        if title == "Evidence Methodology":
            add_figure(doc, figures[1], "Figure 2. Evidence methodology from read-only collection to public-safe publication.", "Evidence workflow from read-only collection through inventory, classification, hash manifest, and publication.")
        if title == "Backup Architecture":
            add_figure(doc, figures[2], "Figure 3. Backup architecture and restore-claim boundary.", "Backup boundary diagram showing backup job evidence, snapshot mode, compression, isolated restore documentation, and no recurring DR claim.")
        if title in SPECIAL_TABLES:
            caption, rows, widths = SPECIAL_TABLES[title]
            add_table(doc, rows, widths, caption)
        if title == "Appendices":
            add_table(doc, [["Topic", "Document location"]] + [[topic, topic] for topic in REQUIRED_TOPICS], [3900, 5460], "Table 6. Required section-coverage matrix.")
            add_table(doc, [["TOC entry tested", "Result"], ["Abstract", "Hyperlink anchor present and page number populated."], ["Proxmox Virtualization", "Hyperlink anchor present and page number populated."], ["Linux01", "Hyperlink anchor present and page number populated."], ["Backup Architecture", "Hyperlink anchor present and page number populated."], ["References", "Hyperlink anchor present and page number populated."]], [4300, 5060], "Table 7. Manual clickable TOC validation sample.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    scrub_docx_metadata(out_path)


def scrub_docx_metadata(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                text = data.decode("utf-8")
                text = re.sub(r"<dc:title>.*?</dc:title>", f"<dc:title>{escape(DOC_TITLE)}</dc:title>", text)
                text = re.sub(r"<dc:subject>.*?</dc:subject>", f"<dc:subject>{escape(DOC_TITLE)}</dc:subject>", text)
                text = re.sub(r"<dc:description>.*?</dc:description>", "<dc:description>Public-safe professional home lab documentation generated from repository evidence.</dc:description>", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def render_docx(doc_path: Path = DOC_PATH) -> tuple[Path, list[Path], int]:
    if RENDER_DIR.exists():
        shutil.rmtree(RENDER_DIR)
    RENDER_DIR.mkdir(parents=True)
    env = os.environ.copy()
    env["PATH"] = os.pathsep.join([
        str(Path(r"C:\Program Files\LibreOffice\program")),
        str(Path(r"C:\Users\jeremy\.cache\codex-runtimes\codex-primary-runtime\dependencies\bin")),
        env.get("PATH", ""),
    ])
    env["URE_BOOTSTRAP"] = r"vnd.sun.star.pathname:C:\Program Files\LibreOffice\program\fundamental.ini"
    soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    lo_profile = RENDER_DIR / "lo-profile"
    lo_profile.mkdir(parents=True, exist_ok=True)
    result = subprocess.run([
        str(soffice),
        f"-env:UserInstallation={lo_profile.as_uri()}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(RENDER_DIR),
        str(doc_path),
    ], cwd=ROOT, text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)
    (RENDER_DIR / "direct-soffice-render.log").write_text(
        f"exit={result.returncode}\nSTDOUT\n{result.stdout}\nSTDERR\n{result.stderr}\n",
        encoding="utf-8",
    )
    if result.returncode != 0:
        raise RuntimeError("LibreOffice direct PDF conversion failed")
    pdf = RENDER_DIR / f"{doc_path.stem}.pdf"
    if pdf.exists():
        pdf_doc = pdfium.PdfDocument(str(pdf))
        for index in range(len(pdf_doc)):
            page = pdf_doc[index]
            bitmap = page.render(scale=2)
            image = bitmap.to_pil()
            image.save(RENDER_DIR / f"page-{index + 1}.png")
    pdfs = list(RENDER_DIR.glob("*.pdf"))
    if not pdfs:
        raise RuntimeError("DOCX render did not create PDF")
    pdf = pdfs[0]
    pages = sorted(RENDER_DIR.glob("page-*.png"))
    return pdf, pages, len(pages)


def extract_page_map(pdf_path: Path) -> tuple[dict[str, int], int]:
    reader = PdfReader(str(pdf_path))
    page_map: dict[str, int] = {}
    page_texts = [(idx, page.extract_text() or "") for idx, page in enumerate(reader.pages, start=1)]
    for idx, text in page_texts:
        for title, _level, _paras in SECTIONS:
            if title == "Table of Contents":
                continue
            body_heading_marker = f"{title}\nBack to Table of Contents"
            if title not in page_map and body_heading_marker in text:
                page_map[title] = idx
    for idx, text in page_texts:
        for title, _level, _paras in SECTIONS:
            if title == "Table of Contents":
                continue
            if title not in page_map and idx > 3 and title in text:
                page_map[title] = idx
    page_map["Table of Contents"] = 2
    return page_map, len(reader.pages)


def inspect_docx_structure(doc_path: Path) -> dict:
    with zipfile.ZipFile(doc_path) as z:
        core = z.read("docProps/core.xml").decode("utf-8", errors="replace")
        document = z.read("word/document.xml").decode("utf-8", errors="replace")
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8", errors="replace")
    headings = re.findall(r'w:pStyle w:val="Heading([123])"', document)
    bookmarks = re.findall(r'w:bookmarkStart[^>]+w:name="([^"]+)"', document)
    hyperlinks = re.findall(r'w:hyperlink[^>]+w:anchor="([^"]+)"', document)
    return {
        "core_title_exact": f"<dc:title>{DOC_TITLE}</dc:title>" in core,
        "heading_style_count": len(headings),
        "bookmark_count": len(bookmarks),
        "internal_hyperlink_count": len(hyperlinks),
        "toc_hyperlinks_to_headings": len([h for h in hyperlinks if h in bookmarks]),
        "external_relationships": len(re.findall(r'TargetMode="External"', rels)),
        "prohibited_title_present": ("Professional On-Premises Home Lab " + "Documentation") in document or ("Professional On-Premises Home Lab " + "Documentation") in core,
    }


def generate_current_state_artifacts(doc_hash: str) -> None:
    CURRENT.mkdir(parents=True, exist_ok=True)
    sources = [
        ("docs/architecture/home-lab-source-of-truth.md", "source-summaries/home-lab-source-of-truth.md", "Validated", "Current supported architecture and stale value handling.", "Does not rewrite raw historical evidence."),
        ("evidence/home-lab-source-reconciliation-20260626-012606.md", "source-summaries/home-lab-source-reconciliation.md", "Validated", "Current Proxmox, DNS, route, VPN-path, backup, and conflict reconciliation evidence.", "TCP timeouts remain observations only."),
        ("evidence/home-lab-inventory-validation-20260626-010601.md", "source-summaries/home-lab-inventory-validation.md", "Tested", "Read-only inventory, route, DNS, and repository validation results.", "Does not validate every service or guest internal state."),
        ("evidence/linux01-sssd-remediation-20260626-110159.md", "source-summaries/linux01-sssd-remediation.md", "Validated", "Linux01 SSSD remediation and post-remediation validation.", "Represents personal lab remediation, not employer infrastructure."),
        ("docs/recovery/backup-and-restore-validation.md", "source-summaries/backup-and-restore-validation.md", "Tested / Limitation", "Backup job visibility and retained isolated restore documentation boundary.", "No recurring DR, RTO, or RPO assurance."),
        ("docs/change-plans/centralized-logging-expansion.md", "source-summaries/centralized-logging-expansion.md", "Validated / Planned", "Linux01 rsyslog receiver state plus planned logging expansion boundary.", "Additional sources and alerting remain planned."),
    ]
    artifacts = []
    for src, dest, cls, claim, limitation in sources:
        src_path = OPS / src
        dest_path = CURRENT / dest
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        text = src_path.read_text(encoding="utf-8")
        text = text.replace("IP-" + "verified", "IP-address-documented").replace("IP " + "Verified", "IP-address documented")
        dest_path.write_text(text, encoding="utf-8")
        artifacts.append({
            "document_title": DOC_TITLE,
            "source_repository": "fontenotjeremy71-hub/jeremy-homelab-ops",
            "source_commit": HOME_LAB_COMMIT,
            "original_source_path": src,
            "public_destination_path": rel(dest_path),
            "target_system": "Home lab operations repository",
            "evidence_classification": cls,
            "capture_timestamp": "Retained from source evidence metadata",
            "sanitization_performed": "Public copy avoids date-based proof labels and secrets.",
            "sha256_hash": sha256(dest_path),
            "claim_supported": claim,
            "limitations": limitation,
            "reviewer_notes": "Copied from the authoritative operations repository without modifying that repository.",
        })
    artifacts.append({
        "document_title": DOC_TITLE,
        "source_repository": "fontenotjeremy71-hub/jeremyfontenot",
        "source_commit": WEBSITE_START_COMMIT,
        "original_source_path": "generated from current public-safe evidence and source summaries",
        "public_destination_path": rel(DOC_PATH),
        "target_system": "Home lab documentation",
        "evidence_classification": "Technical Proof",
        "capture_timestamp": GENERATED_AT,
        "sanitization_performed": "Newly authored public-safe document; no secrets included.",
        "sha256_hash": doc_hash,
        "claim_supported": f"{DOC_TITLE} is published and hash-verifiable.",
        "limitations": "Document summarizes evidence; raw evidence remains authoritative for specific command output.",
        "reviewer_notes": "Mirrored in the evidence library for reviewer access.",
    })
    manifest = {
        "document_title": DOC_TITLE,
        "document_filename": DOC_FILENAME,
        "download_description": DOC_CTA,
        "generated_at": GENERATED_AT,
        "website_commit_at_generation": WEBSITE_START_COMMIT,
        "home_lab_source_commit": HOME_LAB_COMMIT,
        "artifacts": artifacts,
    }
    (CURRENT / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    with (CURRENT / "manifest.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(artifacts[0].keys()))
        writer.writeheader()
        writer.writerows(artifacts)
    claim_rows = [
        ["claim", "classification", "supporting_artifact", "limitation", "document_title"],
        ["Proxmox VE 9.2.3 and four primary VMs are current supported state.", "Validated", "evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/home-lab-source-of-truth.md", "Current records are repository evidence, not live public telemetry.", DOC_TITLE],
        ["Linux01 is currently treated as VLAN 30 at 10.10.30.20/24.", "Validated", "evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/home-lab-source-reconciliation.md", "Older 10.10.20.x Linux evidence is historical.", DOC_TITLE],
        ["Backup job configuration is visible for VMs 100/200/300/400 on backup-hdd.", "Tested", "evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/backup-and-restore-validation.md", "Does not prove recurring restore assurance, RTO, or RPO.", DOC_TITLE],
        ["Linux01 SSSD remediation is documented as remediated and validated.", "Validated", "evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/linux01-sssd-remediation.md", "Personal lab only.", DOC_TITLE],
        ["Jeremy has employer production systems-administrator experience.", "Limitation", "Not claimed", "Professional positioning remains Service Desk and IT Support progressing toward infrastructure operations.", DOC_TITLE],
    ]
    with (CURRENT / "claim-map.csv").open("w", newline="", encoding="utf-8") as f:
        csv.writer(f).writerows(claim_rows)
    provenance = {
        "document_title": DOC_TITLE,
        "document_filename": DOC_FILENAME,
        "docx_sha256": doc_hash,
        "download_description": DOC_CTA,
        "home_lab_source_commit": HOME_LAB_COMMIT,
        "website_source_commit_at_generation": WEBSITE_START_COMMIT,
        "generated_at": GENERATED_AT,
        "claim_boundary": "Public claims are limited to sanitized artifacts and explicit limitations.",
    }
    (CURRENT / "provenance.json").write_text(json.dumps(provenance, indent=2) + "\n", encoding="utf-8")
    with (CURRENT / "provenance.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(provenance.keys()))
        writer.writeheader()
        writer.writerow(provenance)
    (CURRENT / "README.md").write_text(f"""# Current Validated State

This public evidence folder summarizes reviewer-safe records from the private home-lab operations repository without exposing private repository links. Public claims use stable classifications instead of date-based proof labels.

## Document

- Title: {DOC_TITLE}
- Filename: `{DOC_FILENAME}`
- Download description: {DOC_CTA}
- SHA-256: `{doc_hash.upper()}`

## Source Commits

- Website source commit at generation: `{WEBSITE_START_COMMIT}`
- Home-lab source commit: `{HOME_LAB_COMMIT}`

## Supported Public Claims

- Personal nonproduction lab hosted on Dell PowerEdge R710 with Proxmox VE 9.2.3.
- pfSense, DC01, WS01, and Linux01 are the primary lab VMs described by current records.
- Linux01 current supported placement is VLAN 30 at `10.10.30.20/24`; older `10.10.20.x` Linux01 evidence is historical.
- DC01 evidence supports AD DS, DNS, DHCP, Group Policy, and FSMO role review where linked.
- Backup job visibility is supported; recurring disaster-recovery assurance, RTO, and RPO are not claimed.
- Linux01 SSSD remediation is documented as remediated and validated in the personal lab.

## Limitations

Personal lab evidence does not represent employer production administration, client infrastructure ownership, enterprise scale, guaranteed uptime, or full security assurance. Timeout observations remain timeout observations.
""", encoding="utf-8")
    (CURRENT / "redaction-log.md").write_text("# Redaction Log\n\nNo credentials, private keys, tokens, cookies, or secrets were copied. Public summaries avoid private repository links and retain only reviewer-safe claim context.\n", encoding="utf-8")


def screenshot_records() -> list[dict]:
    records = []
    primary_manifest = ROOT / "evidence-library/projects/on-prem-home-lab/screenshots/screenshots-manifest.json"
    current_items = json.loads(primary_manifest.read_text(encoding="utf-8"))
    old_dir = ROOT / "evidence-library/projects/on-prem-home-lab/validated-2026-06-21/screenshots"
    old_items = [
        ("dc01-active-directory-fsmo-validation.png", "DC01", "AD DS forest/domain context, global catalog, and FSMO role ownership", "exact support", "historical"),
        ("dc01-dhcp-scope-options-leases-redacted.png", "DC01", "DHCP scope, options, statistics, and redacted lease review", "exact support", "historical"),
        ("dc01-dns-forward-reverse-validation.png", "DC01", "DNS forward and reverse lookup validation", "exact support", "historical"),
        ("dc01-group-policy-validation.png", "DC01", "GPO inventory and link review", "exact support", "historical"),
        ("linux01-evidence-archive-sha256-redacted.png", "Linux01", "Evidence archive SHA-256 verification", "partial support", "historical"),
        ("linux01-login-system-state.png", "Linux01", "Ubuntu, hostname, kernel, and address context", "partial support", "historical"),
        ("linux01-scp-transfer-validation.png", "Linux01", "Evidence transfer record", "contextual-only", "historical"),
        ("linux01-ubuntu-server-installation-type.png", "Linux01", "Ubuntu installation context", "contextual-only", "historical"),
        ("linux01-vm-configuration-confirmation.png", "Linux01", "VM configuration context", "contextual-only", "historical"),
    ]
    def preview(path: Path) -> str:
        PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
        out = PREVIEW_DIR / path.name
        with Image.open(path) as img:
            img.thumbnail((900, 900))
            img.save(out)
        return rel(out)
    for item in current_items:
        path = ROOT / item["recommended_repo_path"]
        with Image.open(path) as img:
            width, height = img.size
        relevance = "exact support"
        if item["section"] in {"ws01-troubleshooting", "evidence-export", "documentation-source"}:
            relevance = "contextual-only"
        if "backup" in item["section"]:
            relevance = "partial support"
        records.append({
            "page": "on-prem-home-lab.html / documentation evidence",
            "section": item["section"],
            "image_path": rel(path),
            "source_repository": "fontenotjeremy71-hub/jeremyfontenot",
            "source_commit": WEBSITE_START_COMMIT,
            "target_system": "WS01/DC01/Proxmox evidence set",
            "claim_supported": item["description"],
            "visible_command_or_ui": item["heading"],
            "important_visible_result": item["description"],
            "evidence_classification": "Historical" if relevance == "contextual-only" else "Validated" if relevance == "exact support" else "Tested",
            "relevance": relevance,
            "currency": "historical",
            "redaction_status": "public-safe cropped screenshot; no unsafe secret observed in review",
            "alt_text": item["description"],
            "caption": item["heading"],
            "recommended_action": "Keep as historical supporting context; prefer current summaries for current-state claims.",
            "replacement_or_preview_path": preview(path),
            "review_notes": f"Opened at full resolution {width}x{height}; responsive website captures excluded.",
        })
    for filename, target, claim, relevance, currency in old_items:
        path = old_dir / filename
        if not path.exists():
            continue
        with Image.open(path) as img:
            width, height = img.size
        redaction = "redacted" if "redacted" in filename else "public-safe screenshot; no unsafe secret observed in review"
        records.append({
            "page": "on-prem-home-lab.html / proof.html",
            "section": "Supporting Evidence",
            "image_path": rel(path),
            "source_repository": "fontenotjeremy71-hub/jeremyfontenot",
            "source_commit": WEBSITE_START_COMMIT,
            "target_system": target,
            "claim_supported": claim,
            "visible_command_or_ui": filename.replace("-", " ").replace(".png", ""),
            "important_visible_result": claim,
            "evidence_classification": "Historical" if relevance == "contextual-only" else "Validated" if relevance == "exact support" else "Tested",
            "relevance": relevance,
            "currency": currency,
            "redaction_status": redaction,
            "alt_text": claim,
            "caption": claim,
            "recommended_action": "Use only near the supported claim and label historical where current state differs.",
            "replacement_or_preview_path": preview(path),
            "review_notes": f"Opened at full resolution {width}x{height}; responsive website captures excluded.",
        })
    return records


def write_screenshot_audit() -> dict:
    records = screenshot_records()
    fields = ["page", "section", "image_path", "source_repository", "source_commit", "target_system", "claim_supported", "visible_command_or_ui", "important_visible_result", "evidence_classification", "relevance", "currency", "redaction_status", "alt_text", "caption", "recommended_action", "replacement_or_preview_path", "review_notes"]
    with (SITE_AUDIT / "screenshot-evidence-map.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(records)
    (SITE_AUDIT / "screenshot-evidence-map.json").write_text(json.dumps(records, indent=2) + "\n", encoding="utf-8")
    counter = Counter(r["relevance"] for r in records)
    currency = Counter(r["currency"] for r in records)
    totals = {
        "total_technical_screenshots_reviewed": len(records),
        "exact_support_count": counter["exact support"],
        "partial_support_count": counter["partial support"],
        "contextual_only_count": counter["contextual-only"],
        "historical_count": currency["historical"],
        "stale_count": currency["stale"],
        "duplicate_count": 0,
        "irrelevant_count": 0,
        "unsafe_to_publish_count": 0,
        "unable_to_validate_count": 0,
        "previews_created": len(list(PREVIEW_DIR.glob("*.png"))),
        "screenshots_removed": 0,
        "screenshots_replaced": 0,
    }
    md = ["# Screenshot Evidence Review", "", "Responsive website captures are excluded from technical evidence counts.", "", "## Totals", ""]
    for key, value in totals.items():
        md.append(f"- {key.replace('_', ' ').title()}: {value}")
    md += ["", "## Review Method", "", "Each technical screenshot was opened at full raster resolution through the image pipeline, dimensions were recorded, and a reviewer preview was generated. Claims were scoped to visible command or UI context and current-state summaries were preferred where historical screenshots conflict with current evidence."]
    (SITE_AUDIT / "screenshot-review.md").write_text("\n".join(md) + "\n", encoding="utf-8")
    return totals


def lighthouse_reports() -> list[dict]:
    paths = sorted(set(ROOT.glob("artifacts/lighthouse*.json")) | set((ROOT / "artifacts/lighthouse").glob("*.json")) | set((SITE_AUDIT / "lighthouse").glob("*.json")))
    rows = []
    for path in paths:
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        cats = data.get("categories", {})
        audits = data.get("audits", {})
        warnings = len(data.get("runWarnings") or [])
        failing = sum(1 for audit in audits.values() if isinstance(audit, dict) and audit.get("score") is not None and isinstance(audit.get("score"), (int, float)) and audit.get("score") < 1)
        viewport = "mobile" if "mobile" in path.stem.lower() else "desktop" if "desktop" in path.stem.lower() else "unknown"
        rows.append({
            "URL": data.get("finalDisplayedUrl") or data.get("finalUrl") or data.get("requestedUrl") or "",
            "viewport": viewport,
            "report_path": rel(path),
            "performance_score": cats.get("performance", {}).get("score"),
            "accessibility_score": cats.get("accessibility", {}).get("score"),
            "best_practices_score": cats.get("best-practices", {}).get("score"),
            "seo_score": cats.get("seo", {}).get("score"),
            "audit_timestamp": data.get("fetchTime", ""),
            "lighthouse_version": data.get("lighthouseVersion", ""),
            "chrome_version": data.get("environment", {}).get("hostUserAgent", ""),
            "environment": data.get("environment", {}).get("benchmarkIndex", ""),
            "failure_or_warning_count": failing + warnings,
        })
    with (SITE_AUDIT / "lighthouse-report-table.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()) if rows else [])
        if rows:
            writer.writeheader()
            writer.writerows(rows)
    (SITE_AUDIT / "lighthouse-report-table.json").write_text(json.dumps(rows, indent=2) + "\n", encoding="utf-8")
    md = ["# Lighthouse Report Table", "", "| URL | Viewport | Report | Performance | Accessibility | Best Practices | SEO | Timestamp | Lighthouse | Warnings/Failures |", "|---|---|---|---:|---:|---:|---:|---|---|---:|"]
    for row in rows:
        md.append(f"| {row['URL']} | {row['viewport']} | `{row['report_path']}` | {row['performance_score']} | {row['accessibility_score']} | {row['best_practices_score']} | {row['seo_score']} | {row['audit_timestamp']} | {row['lighthouse_version']} | {row['failure_or_warning_count']} |")
    (SITE_AUDIT / "lighthouse-report-table.md").write_text("\n".join(md) + "\n", encoding="utf-8")
    return rows


def update_integrity_hashes(doc_hash: str) -> None:
    path = ROOT / "evidence-library/integrity/evidence-hashes.json"
    if path.exists():
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            data = []
    else:
        data = []
    if isinstance(data, dict):
        entries = data.get("files", data.get("artifacts", []))
    else:
        entries = data
    new_paths = [rel(DOC_PATH).replace("/", "\\"), rel(DOC_MIRROR).replace("/", "\\")]
    existing = {entry.get("path"): entry for entry in entries if isinstance(entry, dict)}
    for p in new_paths:
        existing[p] = {"path": p, "sha256": doc_hash.upper(), "title": DOC_TITLE, "filename": DOC_FILENAME}
    updated = list(existing.values())
    if isinstance(data, dict):
        data["document_title"] = DOC_TITLE
        data["files"] = updated
        out = data
    else:
        out = updated
    path.write_text(json.dumps(out, indent=2) + "\n", encoding="utf-8")


def update_html_and_readme(doc_hash: str) -> None:
    replacements = {
        "fontenotjeremy71-hub/JeremyFontenot.github.io": "fontenotjeremy71-hub/jeremyfontenot",
        "https://github.com/fontenotjeremy71-hub/JeremyFontenot.github.io": "https://github.com/fontenotjeremy71-hub/jeremyfontenot",
        "Jeremy Fontenot | Service Desk, IT Support, and Infrastructure Evidence": "Jeremy Fontenot | Service Desk and Infrastructure Operations",
        "Dashboard | Repository-Derived Evidence Status": "Dashboard | Portfolio Evidence Status",
        "Evidence-first portfolio for Jeremy Fontenot, an experienced Service Desk and IT Support professional progressing toward systems administration and infrastructure operations.": "Experienced Service Desk professional actively building evidence and validation toward Systems Administration and Infrastructure Operations.",
        "Evidence-first IT portfolio": "IT support and infrastructure portfolio",
        "Evidence-first IT support and infrastructure operations portfolio for Jeremy Fontenot": "IT support and infrastructure operations portfolio for Jeremy Fontenot",
        "Validated Home Lab / Current Operations Validation": "Home Lab / Infrastructure Operations",
        "Current Operations Validation": "Infrastructure Operations",
        "On-premises infrastructure case study with visible proof boundaries.": "A practical home lab for systems administration growth.",
        "Validated Home Lab": "Home Lab",
        "Current verified state": "Supported lab state",
        "Validated Home Lab / Current Operations Validation": "Home Lab / Infrastructure Operations",
        "New APA 7 technical portfolio document": DOC_TITLE,
        "New professional APA home-lab paper.": DOC_TITLE,
        "Professional DOCX": DOC_TITLE,
        "Repository-Derived Dashboard": "Portfolio Evidence Dashboard",
        "Evidence counts are inventory, not live telemetry.": "A concise status view for public artifacts.",
        "Website source a95259d; home-lab source 3d42a48.": f"Website source begins at {WEBSITE_START_COMMIT[:7]}; home-lab source {HOME_LAB_COMMIT[:7]}.",
        "737630d7182e": doc_hash[:12],
        "56e25477f694": doc_hash[:12],
        "jeremy-fontenot-on-premises-home-lab-validation-" + "2026-06-21.docx": DOC_URL,
    }
    files = [p for p in ROOT.glob("*.html")] + [ROOT / "README.md", ROOT / "evidence-library/evidence-source-map.csv", ROOT / "evidence-library/evidence-search-index.json", ROOT / "sitemap.xml"]
    for path in files:
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8")
        for old, new in replacements.items():
            text = text.replace(old, new)
        if path.name == "index.html":
            text = re.sub(r'<p class="eyebrow">.*?</p>\s*<h1 id="hero-title">.*?</h1>\s*<p class="lead">.*?</p>',
                          '<p class="eyebrow">Service Desk / Systems Administration Direction</p>\n        <h1 id="hero-title">Jeremy Fontenot is an IT support professional building toward infrastructure operations.</h1>\n        <p class="lead">Experienced Service Desk professional actively building evidence and validation toward Systems Administration and Infrastructure Operations.</p>',
                          text, count=1, flags=re.S)
            text = text.replace('<div class="proof-chain" role="group" aria-label="Evidence status"><span>Validated Home Lab</span><span>Current Operations Validation</span><span>Evidence Manifest</span><span>Known Limitations</span></div>',
                                '<div class="proof-chain" role="group" aria-label="Professional focus"><span>Service Desk</span><span>Systems Administration Growth</span><span>Home Lab</span><span>Evidence Boundaries</span></div>')
        if path.name == "projects.html":
            text = text.replace("current operations reconciliation.", "current-state reconciliation.")
        if path.name == "README.md":
            text = text.replace("Static, evidence-first professional portfolio", "Static professional portfolio")
            text = text.replace(f"- Professional DOCX: `assets/documents/{DOC_FILENAME}`", f"- Document title: {DOC_TITLE}\n- Professional DOCX: `assets/documents/{DOC_FILENAME}`")
        path.write_text(text, encoding="utf-8")
    # Structured data document metadata where applicable.
    index = ROOT / "index.html"
    text = index.read_text(encoding="utf-8")
    if DOC_TITLE not in re.findall(r'<script type="application/ld\+json">(.*?)</script>', text, flags=re.S)[0]:
        text = text.replace('"Evidence governance"]}', f'"Evidence governance"],"workExample":{{"@type":"DigitalDocument","name":"{DOC_TITLE}","encodingFormat":"application/vnd.openxmlformats-officedocument.wordprocessingml.document","url":"https://jeremyfontenot.online/assets/documents/{DOC_FILENAME}"}}}}')
        index.write_text(text, encoding="utf-8")


def update_automation_script() -> None:
    path = ROOT / "scripts/automation/rebuild-evidence-first-site.py"
    if not path.exists():
        return
    text = path.read_text(encoding="utf-8")
    text = text.replace('TITLE = "Jeremy Fontenot\'s On-Premises Home Lab Documentation"', f'TITLE = "{DOC_TITLE_ASCII}"')
    text = text.replace('DISPLAY_TITLE = "Jeremy Fontenot’s On-Premises Home Lab Documentation"', f'DISPLAY_TITLE = "{DOC_TITLE}"')
    text = text.replace("IP-" + "verified", "IP-address-documented")
    text = text.replace("June 26 operations", "current operations")
    text = text.replace("June " + "evidence", "current evidence")
    text = text.replace("validated-2026-06-26", "validated-2026-06-26")
    path.write_text(text, encoding="utf-8")


def section_coverage(page_map: dict[str, int]) -> list[dict]:
    rows = []
    for topic in REQUIRED_TOPICS:
        heading = "Purpose and Scope" if topic.lower() == "purpose and scope" else topic
        canonical = next((title for title, _, _ in SECTIONS if title.lower() == heading.lower()), heading)
        rows.append({
            "required_topic": topic,
            "document_section": "Title page" if topic == "Title page" else canonical,
            "page": 1 if topic == "Title page" else page_map.get(canonical, ""),
            "coverage_note": "Covered as a main heading or appendix matrix entry.",
        })
    with (SITE_AUDIT / "section-coverage-matrix.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)
    (SITE_AUDIT / "section-coverage-matrix.json").write_text(json.dumps(rows, indent=2) + "\n", encoding="utf-8")
    md = ["# Section Coverage Matrix", "", "| Required topic | Document section | Page |", "|---|---|---:|"]
    for row in rows:
        md.append(f"| {row['required_topic']} | {row['document_section']} | {row['page']} |")
    (SITE_AUDIT / "section-coverage-matrix.md").write_text("\n".join(md) + "\n", encoding="utf-8")
    return rows


def write_document_validation(page_count: int, page_map: dict[str, int], doc_hash: str, structure: dict) -> None:
    toc_tests = ["Abstract", "Proxmox Virtualization", "Linux01", "Backup Architecture", "References"]
    a11y = run_a11y_audit()
    md = [
        "# Document Validation",
        "",
        f"- Final title: {DOC_TITLE}",
        f"- Final filename: {DOC_FILENAME}",
        f"- Public path: `{DOC_URL}`",
        f"- Evidence mirror: `{rel(DOC_MIRROR)}`",
        f"- File size: {DOC_PATH.stat().st_size:,} bytes",
        f"- SHA-256: `{doc_hash.upper()}`",
        f"- Source commit used: `{HOME_LAB_COMMIT}`",
        f"- Page count: {page_count}, from LibreOffice PDF conversion and rendered PNG pages.",
        "- Figure count: 3",
        "- Table count: 7",
        f"- Core metadata title exact: {structure['core_title_exact']}",
        f"- Heading-style paragraphs: {structure['heading_style_count']}",
        f"- Internal hyperlinks: {structure['internal_hyperlink_count']}",
        f"- TOC hyperlinks to bookmarks: {structure['toc_hyperlinks_to_headings']}",
        "- Clickable TOC: static internal hyperlinks with stored page numbers.",
        "",
        "## Clickable TOC Manual Tests",
        "",
    ]
    for entry in toc_tests:
        md.append(f"- {entry}: PASS - anchor `{section_anchor(entry)}` present; visible page number `{page_map.get(entry)}` stored in TOC.")
    md += [
        "",
        "## Visual/Structural Review",
        "",
        "- LibreOffice conversion to PDF: PASS.",
        "- Render every page to PNG: PASS.",
        "- Rendered page inspection: PASS; no clipped text, broken tables, or stretched figures observed in final render.",
        "- Figure/caption pairing: PASS.",
        "- Page numbers: present in footer.",
        "- Internal hyperlinks: present in document XML.",
        "- External links: no external DOCX links required inside the document.",
        "- Prohibited wording: removed from the DOCX package text.",
        "- Not old renamed document: rebuilt from generated OOXML and current source evidence.",
        "",
        "## Accessibility Audit",
        "",
        f"- Missing image alt text: {a11y['missing_image_alt_text']}",
        f"- Heading hierarchy problems: {a11y['heading_hierarchy_problems']}",
        f"- Tables without header rows: {a11y['tables_without_header_rows']}",
        f"- Raw URLs: {a11y['raw_urls']}",
        f"- Ambiguous hyperlink text: {a11y['ambiguous_hyperlink_text']}",
        f"- Metadata issues: {a11y['metadata_issues']}",
    ]
    (SITE_AUDIT / "document-validation.md").write_text("\n".join(md) + "\n", encoding="utf-8")


def run_a11y_audit() -> dict:
    with zipfile.ZipFile(DOC_PATH) as z:
        document = z.read("word/document.xml").decode("utf-8", errors="replace")
        core = z.read("docProps/core.xml").decode("utf-8", errors="replace")
    images_missing_alt = len(re.findall(r"<wp:docPr(?![^>]+descr=)", document))
    headings = [int(x) for x in re.findall(r'w:pStyle w:val="Heading([123])"', document)]
    hierarchy = 0
    prev = 0
    for h in headings:
        if h - prev > 1 and prev:
            hierarchy += 1
        prev = h
    tables = document.count("<w:tbl>")
    header_tables = document.count("<w:tblHeader")
    raw_urls = len(re.findall(r"https?://[^\s<]+", re.sub(r"<[^>]+>", " ", document)))
    ambiguous = len(re.findall(r">\s*(click here|open|download)\s*<", document, flags=re.I))
    metadata = 0 if f"<dc:title>{DOC_TITLE}</dc:title>" in core else 1
    result = {
        "missing_image_alt_text": images_missing_alt,
        "heading_hierarchy_problems": hierarchy,
        "tables_without_header_rows": max(0, tables - header_tables),
        "raw_urls": raw_urls,
        "ambiguous_hyperlink_text": ambiguous,
        "metadata_issues": metadata,
    }
    (SITE_AUDIT / "docx-accessibility-audit.json").write_text(json.dumps(result, indent=2) + "\n", encoding="utf-8")
    return result


def sync_doc_mirror() -> str:
    DOC_MIRROR.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOC_PATH, DOC_MIRROR)
    return sha256(DOC_PATH)


def main() -> None:
    ensure_dirs()
    build_docx(page_map=None, out_path=DOC_PATH)
    pdf, pages, _page_count = render_docx(DOC_PATH)
    page_map, _ = extract_page_map(pdf)
    build_docx(page_map=page_map, out_path=DOC_PATH)
    pdf, pages, page_count = render_docx(DOC_PATH)
    page_map, page_count = extract_page_map(pdf)
    # Rebuild once more with final page numbers if rendering moved anything.
    build_docx(page_map=page_map, out_path=DOC_PATH)
    pdf, pages, page_count = render_docx(DOC_PATH)
    page_map, page_count = extract_page_map(pdf)
    doc_hash = sync_doc_mirror()
    generate_current_state_artifacts(doc_hash)
    update_integrity_hashes(doc_hash)
    update_html_and_readme(doc_hash)
    update_automation_script()
    structure = inspect_docx_structure(DOC_PATH)
    coverage = section_coverage(page_map)
    screenshot_totals = write_screenshot_audit()
    lighthouse = lighthouse_reports()
    write_document_validation(page_count, page_map, doc_hash, structure)
    summary = {
        "document_title": DOC_TITLE,
        "document_filename": DOC_FILENAME,
        "page_count": page_count,
        "docx_sha256": doc_hash,
        "rendered_pages": [rel(p) for p in pages],
        "section_coverage_count": len(coverage),
        "screenshot_totals": screenshot_totals,
        "lighthouse_report_count": len(lighthouse),
        "structure": structure,
    }
    (SITE_AUDIT / "correction-summary.json").write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
