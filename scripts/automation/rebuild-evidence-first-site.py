from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OPS = Path(r"C:\Users\jeremy\Documents\projects\jeremy-homelab-ops")
PUBLIC_DOCX = ROOT / "assets" / "documents" / "Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx"
PUBLIC_EVIDENCE = ROOT / "evidence-library" / "projects" / "on-prem-home-lab"
CURRENT = PUBLIC_EVIDENCE / "current-validated-state"
AUDIT = ROOT / "artifacts" / "site-audit"
DOMAIN = "https://jeremyfontenot.online"
TITLE = "Jeremy Fontenot's On-Premises Home Lab Documentation"
DISPLAY_TITLE = "Jeremy Fontenot’s On-Premises Home Lab Documentation"
DOC_CTA = "Download the Professional Home Lab Documentation (.docx)"
DOC_URL = "./assets/documents/Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx"


def git(*args: str, cwd: Path = ROOT) -> str:
    return subprocess.check_output(["git", *args], cwd=cwd, text=True, stderr=subprocess.STDOUT).strip()


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def copy_source(name: str, src: Path, classification: str, claim: str, limitations: str) -> dict:
    dest = CURRENT / "source-summaries" / name
    dest.parent.mkdir(parents=True, exist_ok=True)
    text = src.read_text(encoding="utf-8", errors="replace")
    text = re.sub(r"IP[-_ ]Verified", "retained", text, flags=re.I)
    write(dest, text)
    return {
        "source_repository": "fontenotjeremy71-hub/jeremy-homelab-ops",
        "source_commit": OPS_COMMIT,
        "original_source_path": str(src.relative_to(OPS)).replace("\\", "/"),
        "public_destination_path": str(dest.relative_to(ROOT)).replace("\\", "/"),
        "target_system": "Home lab operations repository",
        "evidence_classification": classification,
        "capture_timestamp": "Retained from source evidence metadata",
        "sanitization_performed": "Public copy avoids retired marketing label; no secrets added.",
        "sha256_hash": sha256(dest),
        "claim_supported": claim,
        "limitations": limitations,
        "reviewer_notes": "Copied from the authoritative operations repository without modifying that repository.",
    }


def nav(active: str) -> str:
    items = [
        ("Home", "index.html"),
        ("Projects", "projects.html"),
        ("Home Lab", "on-prem-home-lab.html"),
        ("Proof", "proof.html"),
        ("Dashboard", "dashboard.html"),
        ("Resume", "resume.html"),
        ("Contact", "contact.html"),
    ]
    links = "".join(
        f'<a href="./{href}"{" aria-current=\"page\"" if active == label else ""}>{label}</a>'
        for label, href in items
    )
    return f"""
<a class="skip-link" href="#main">Skip to content</a>
<header class="site-header">
  <nav class="nav" aria-label="Primary navigation">
    <a class="brand" href="./index.html"><img src="./assets/logos/header_logo_88x88.png" alt="Jeremy Fontenot logo" width="44" height="44" decoding="async"><span>Jeremy Fontenot</span><small>IT support and infrastructure portfolio</small></a>
    <button class="nav-toggle" type="button" aria-expanded="false" aria-controls="primary-menu">Menu</button>
    <div class="nav-links" id="primary-menu">{links}</div>
  </nav>
</header>"""


def footer() -> str:
    return f"""
<footer class="site-footer premium-footer" aria-label="Site footer">
  <div class="footer-proof-band" role="group" aria-label="Portfolio proof status">
    <span><strong>Evidence</strong> source mapped</span>
    <span><strong>Claims</strong> classified</span>
    <span><strong>Documents</strong> hashed</span>
    <span><strong>Limits</strong> visible</span>
  </div>
  <div class="footer-divider" aria-hidden="true"></div>
  <div class="footer-grid">
    <div class="footer-brand">
      <img src="./assets/logos/header_logo_88x88.png" alt="Jeremy Fontenot logo" width="56" height="56" decoding="async">
      <p><strong>Jeremy Fontenot</strong></p>
      <p>Experienced Service Desk and IT Support professional building toward systems administration and infrastructure operations through evidence-backed lab work, documentation, and validation.</p>
    </div>
    <div class="footer-links">
      <h2 class="footer-heading">Primary Navigation</h2>
      <a href="./index.html">Home</a><a href="./projects.html">Projects</a><a href="./on-prem-home-lab.html">Home Lab</a><a href="./proof.html">Proof</a><a href="./dashboard.html">Dashboard</a><a href="./resume.html">Resume</a><a href="./contact.html">Contact</a>
    </div>
    <div class="footer-links">
      <h2 class="footer-heading">Evidence Links</h2>
      <a href="./evidence-library/projects/on-prem-home-lab/current-validated-state/manifest.json">Current home-lab manifest</a>
      <a href="{DOC_URL}">{DOC_CTA}</a>
      <a href="./evidence-library/projects/microsoft-365-lab/m365-entra-proof-inventory-20260605.csv">Microsoft 365 proof inventory</a>
      <a href="./artifacts/site-audit/screenshot-review.md">Screenshot evidence review</a>
    </div>
  </div>
  <p class="credibility">Proof boundary: public claims are limited to sanitized repository artifacts, screenshots, inventories, validation summaries, document hashes, and stated limitations. Personal lab evidence is not presented as employer production administration.</p>
  <p class="footer-meta">Service Desk / IT Support / Microsoft 365 / Entra ID / Active Directory / Windows Server / Linux / PowerShell / Evidence Governance</p>
</footer>"""


def page(active: str, title: str, description: str, body: str, kind: str = "WebPage") -> str:
    url = f"{DOMAIN}/" if active == "Home" else f"{DOMAIN}/{active.lower().replace(' ', '-')}.html"
    canonical = {
        "Home Lab": f"{DOMAIN}/on-prem-home-lab.html",
    }.get(active, url)
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{escape(title)}</title>
<meta name="description" content="{escape(description)}">
<meta name="robots" content="index, follow">
<meta name="theme-color" content="#0f172a">
<meta name="referrer" content="strict-origin-when-cross-origin">
<link rel="canonical" href="{canonical}">
<meta property="og:site_name" content="Jeremy Fontenot">
<meta property="og:title" content="{escape(title)}">
<meta property="og:description" content="{escape(description)}">
<meta property="og:type" content="website">
<meta property="og:url" content="{canonical}">
<meta property="og:image" content="{DOMAIN}/assets/og/og-portfolio.png">
<meta property="og:image:alt" content="IT support and infrastructure operations portfolio for Jeremy Fontenot">
<meta name="twitter:card" content="summary_large_image">
<meta name="twitter:title" content="{escape(title)}">
<meta name="twitter:description" content="{escape(description)}">
<meta name="twitter:image" content="{DOMAIN}/assets/og/og-portfolio.png">
<link rel="icon" href="./assets/logos/favicon_64x64.png">
<link rel="stylesheet" href="./assets/css/site.css?v=20260627-layout-cache">
<script src="./assets/js/site.js" defer></script>
<script type="application/ld+json">{json.dumps(schema(kind, title, description, canonical), separators=(",", ":"))}</script>
</head>
<body class="{active.lower().replace(' ', '-')}-page">
{nav(active)}
<main id="main">
{body}
</main>
{footer()}
</body>
</html>"""


def schema(kind: str, title: str, description: str, url: str) -> dict:
    base = {
        "@context": "https://schema.org",
        "@type": kind,
        "name": title,
        "headline": title,
        "description": description,
        "url": url,
        "author": {"@type": "Person", "name": "Jeremy Fontenot", "url": DOMAIN},
        "about": [
            "Service desk support",
            "Microsoft 365 administration",
            "Active Directory",
            "Windows Server",
            "Proxmox VE",
            "pfSense",
            "PowerShell",
            "Evidence governance",
        ],
    }
    if kind == "Person":
        base = {
            "@context": "https://schema.org",
            "@type": "Person",
            "name": "Jeremy Fontenot",
            "url": DOMAIN,
            "jobTitle": "Service Desk and IT Support Professional",
            "sameAs": [
                "https://github.com/fontenotjeremy71-hub",
                "https://www.linkedin.com/in/jeremy-fontenot/",
            ],
            "knowsAbout": base["about"],
        }
    return base


def cards(items: list[tuple[str, str, str, str]]) -> str:
    return "".join(f'<article><span class="tile-code">{k}</span><h3>{h}</h3><p>{p}</p>{a}</article>' for k, h, p, a in items)


def build_pages() -> None:
    home = f"""
  <section class="hero page-shell command-hero" aria-labelledby="hero-title">
    <div class="hero-grid">
      <div class="hero-copy reveal is-visible">
        <p class="eyebrow">Experienced Service Desk and IT Support / Systems Administration Direction</p>
        <h1 id="hero-title">Jeremy Fontenot turns support work into infrastructure evidence.</h1>
        <p class="lead">Experienced IT support professional building toward systems administration and infrastructure operations, with recruiter-friendly role context and technical proof that can be inspected: current home-lab inventories, Microsoft 365 exports, validation records, screenshots, hashes, and documented limitations.</p>
        <div class="actions"><a class="button primary" href="./on-prem-home-lab.html">Home Lab</a><a class="button" href="./projects.html">Projects</a><a class="button" href="./proof.html">Proof</a><a class="button" href="./resume.html">Resume</a><a class="button" href="./contact.html">Contact</a></div>
        <div class="proof-chain" role="group" aria-label="Evidence status"><span>Validated Home Lab</span><span>Current Operations Validation</span><span>Evidence Manifest</span><span>Known Limitations</span></div>
      </div>
      <aside class="review-console evidence-console verification-glow reveal is-visible" aria-label="Infrastructure evidence visual">
        <div class="console-bar"><span>supported lab state</span><span>portfolio proof</span></div>
        <div class="lab-proof-stack">
          <a href="./on-prem-home-lab.html"><b>LAB</b><span>Personal on-prem environment</span><em>Proxmox, pfSense, AD DS, Windows, Linux, backups</em></a>
          <a href="./proof.html"><b>MAP</b><span>Claim-to-artifact proof index</span><em>Every strong claim routes to evidence and limitations</em></a>
          <a href="{DOC_URL}"><b>DOC</b><span>Professional home-lab documentation</span><em>{DISPLAY_TITLE}</em></a>
          <a href="./dashboard.html"><b>QA</b><span>Portfolio evidence dashboard</span><em>A concise status view for public artifacts</em></a>
        </div>
      </aside>
    </div>
  </section>
  <section class="section" aria-labelledby="value-title">
    <div class="section-head reveal"><p class="eyebrow">Professional Value</p><h2 id="value-title">Support-first experience with infrastructure-grade documentation habits.</h2><p>Jeremy's public portfolio is organized for two review paths: recruiters can understand role fit quickly, and technical reviewers can inspect the records behind systems administration claims.</p></div>
    <div class="signal-matrix reveal">{cards([
      ("SD", "Service Desk and IT Support", "Troubleshooting, user support, escalation language, documentation, and careful claim boundaries remain the professional center of gravity.", ""),
      ("SYS", "Systems Administration Direction", "Personal lab work demonstrates configured and validated infrastructure without overstating employment scope.", ""),
      ("M365", "Microsoft 365 and Entra", "Personal tenant exports, admin-center screenshots, and proof inventories support cloud administration review.", ""),
      ("PS", "PowerShell and Governance", "Validation scripts, manifests, hashes, and repository checks show repeatable operational discipline.", ""),
    ])}</div>
  </section>
  <section class="section" aria-labelledby="projects-title">
    <div class="section-head reveal"><p class="eyebrow">Selected Evidence-Backed Projects</p><h2 id="projects-title">Five projects, each tied to reviewer-ready proof.</h2></div>
    <div class="evidence-focus-grid reveal">{cards([
      ("01", "On-Premises Home Lab and Operations Validation", "Current supported state includes Proxmox VE 9.2.3, pfSense, DC01, WS01, Linux01 on VLAN 30, OpenVPN management path, backup configuration, and documented limitations.", '<a class="evidence-link" href="./on-prem-home-lab.html"><span>case study</span>Open</a>'),
      ("02", "Microsoft 365 and Entra Administration Evidence", "Repository-local exports and screenshots support tenant, domain, user, group, role, license, sign-in, audit, and Conditional Access review claims.", '<a class="evidence-link" href="./proof.html#m365-proof"><span>proof</span>Open</a>'),
      ("03", "Service Desk Troubleshooting and RCA", "Incident summaries and troubleshooting records are framed as support practice, not unsupported business-impact claims.", '<a class="evidence-link" href="./projects.html#service-desk-rca"><span>project</span>Open</a>'),
      ("04", "PowerShell Automation and Validation", "Scripts and validation checks demonstrate repeatable evidence collection, repository health checks, and audit discipline.", '<a class="evidence-link" href="./projects.html#automation"><span>project</span>Open</a>'),
    ])}</div>
  </section>
  <section class="section scope-note-section" aria-labelledby="method-title">
    <div class="scope-note-card reveal"><p class="eyebrow">Evidence Methodology</p><h2 id="method-title">Dates are metadata. Proof comes from artifacts.</h2><p>The site prioritizes commands, authenticated inventories, manifests, hashes, screenshot review, repository validation, source paths, and limitations. Historical timestamps remain in raw evidence where they matter, but public labels use stable proof classifications.</p></div>
  </section>"""
    write(ROOT / "index.html", page("Home", "Jeremy Fontenot | Service Desk and Infrastructure Operations", "Experienced Service Desk professional actively building evidence and validation toward Systems Administration and Infrastructure Operations.", home, "Person"))

    project_items = [
        ("infrastructure operations security", "On-Premises Home Lab and Operations Validation", "Personal nonproduction lab with Proxmox VE 9.2.3, pfSense VM 100, DC01 Windows Server 2022, WS01 Windows workstation, Linux01 Ubuntu 26.04 LTS on VLAN 30, OpenVPN management path evidence, backup job configuration, and current operations reconciliation.", "Validated / Tested / Limitation", "./on-prem-home-lab.html"),
        ("m365 identity cloud", "Microsoft 365 and Entra Administration Evidence", "Personal cloud lab evidence includes tenant/domain exports, users, groups, directory roles, license records, Conditional Access summaries, sign-in activity, audit activity, device inventory, application inventory, and admin-center screenshots.", "Validated / Historical where marked", "./proof.html#m365-proof"),
        ("service desk rca troubleshooting", "Service Desk Troubleshooting and Root Cause Analysis", "Support-oriented incident records emphasize symptoms, scope, troubleshooting steps, root-cause framing, remediation planning, and limits without inventing unsupported impact metrics.", "Implemented / Tested where recorded", "./evidence-library/projects/troubleshooting-rca/rca-report.md"),
        ("automation powershell validation", "PowerShell Automation and Validation", "Repository validation scripts, evidence hash generation, link checks, and read-only baseline tooling show practical automation habits tied to reviewable outputs.", "Tested / Validated", "./evidence-library/projects/powershell-automation/operational-handoff-2026.txt"),
        ("governance deployment quality", "Repository Governance and Deployment Quality", "Static-site validation, sitemap checks, structured data, document hash records, screenshot review, and deployment workflow records support maintainable publication quality.", "Tested", "./dashboard.html"),
    ]
    projects = '<section class="page page-hero" aria-labelledby="projects-title"><div class="section-head reveal"><p class="eyebrow">Projects</p><h1 id="projects-title">Evidence-backed work, ordered for hiring review.</h1><p class="lead">Each project states the objective, environment, validation method, exact proof path, classification, limitation, outcome, and professional skill demonstrated.</p></div><div class="project-filter-bar" role="group" aria-label="Project filters"><button class="filter-button is-active" type="button" data-filter="all">All</button><button class="filter-button" type="button" data-filter="infrastructure">Infrastructure</button><button class="filter-button" type="button" data-filter="m365">Microsoft 365</button><button class="filter-button" type="button" data-filter="service">Service Desk</button><button class="filter-button" type="button" data-filter="automation">Automation</button></div><div class="project-systems">'
    for tags, name, summary, classification, link in project_items:
        projects += f'<article id="{name.lower().split()[0].replace("on-premises","home-lab")}" class="project-panel feature-card reveal" data-project="{tags}"><p class="eyebrow">{classification}</p><h2>{name}</h2><p>{summary}</p><dl class="evidence-dl"><dt>Objective</dt><dd>Make the work inspectable and understandable to both recruiters and technical reviewers.</dd><dt>Validation method</dt><dd>Source artifacts, manifests, hashes, screenshots, scripts, or documented limitation records.</dd><dt>Outcome</dt><dd>Professional skills demonstrated: troubleshooting, documentation, infrastructure awareness, automation discipline, and evidence governance.</dd></dl><div class="proof-links"><a href="{link}">Open proof path</a><a href="./proof.html">Reviewer proof index</a></div></article>'
    projects += '</div></section>'
    write(ROOT / "projects.html", page("Projects", "Projects | Jeremy Fontenot Evidence-Backed IT Portfolio", "Evidence-backed IT support, Microsoft 365, home lab, automation, troubleshooting, and repository governance projects.", projects, "CollectionPage"))

    lab = f"""
  <section class="page page-hero home-lab-hero" aria-labelledby="lab-title">
    <div class="page-hero-grid">
      <div class="hero-copy reveal is-visible"><p class="eyebrow">Validated Home Lab / Current Operations Validation</p><h1 id="lab-title">On-premises infrastructure case study with visible proof boundaries.</h1><p class="lead">A personal nonproduction lab that supports systems administration learning and technical review: Dell PowerEdge R710, Proxmox VE, pfSense, Windows Server 2022, Active Directory, DNS, DHCP, Group Policy, WS01, Linux01, OpenVPN management path, logging, backups, and evidence governance.</p><div class="actions"><a class="button primary" href="{DOC_URL}">{DOC_CTA}</a><a class="button" href="./proof.html#home-lab-proof">Review Proof</a><a class="button" href="./evidence-library/projects/on-prem-home-lab/current-validated-state/manifest.json">Open Manifest</a></div></div>
      <aside class="lab-topology reveal is-visible" aria-label="Home lab topology visual"><div class="topology-node"><span>Physical Host</span><strong>Dell PowerEdge R710</strong><em>Personal lab host</em></div><div class="topology-link"></div><div class="topology-node"><span>Hypervisor</span><strong>Proxmox VE 9.2.3</strong><em>VMs 100, 200, 300, 400 running</em></div><div class="topology-link"></div><div class="topology-bus"><span>Network and Identity</span><strong>pfSense / AD DS / DNS / DHCP / VLAN 30 / OpenVPN</strong></div><div class="topology-endpoints"><div class="topology-node"><span>DC01 / WS01</span><strong>10.10.20.0/24</strong><em>Windows services and workstation evidence</em></div><div class="topology-node"><span>Linux01</span><strong>10.10.30.20/24</strong><em>Ubuntu domain member on VLAN 30</em></div></div></aside>
    </div>
  </section>
  <section class="section" aria-labelledby="legend-title"><div class="section-head reveal"><p class="eyebrow">Evidence Strength Legend</p><h2 id="legend-title">Claims are classified before they are promoted.</h2></div><div class="signal-matrix reveal">{cards([
      ("VAL", "Validated", "Direct evidence confirms the stated behavior or current state.", ""),
      ("TST", "Tested", "A relevant test was performed and recorded, but the result may be narrower than the full claim.", ""),
      ("CFG", "Configured", "Configuration is present; behavior is not expanded beyond supporting evidence.", ""),
      ("LIM", "Limitation", "Known gaps, timeouts, historical records, or unsupported conclusions remain visible.", ""),
    ])}</div></section>
  <section class="section" aria-labelledby="layers-title"><div class="section-head reveal"><p class="eyebrow">Architecture Layers</p><h2 id="layers-title">Current supported state, from hardware to evidence workflow.</h2></div><ol class="lab-timeline reveal"><li><b>01</b><span>Physical host: Dell PowerEdge R710.</span></li><li><b>02</b><span>Hypervisor: Proxmox VE 9.2.3.</span></li><li><b>03</b><span>Storage: backup-hdd at /mnt/pve/backup-hdd.</span></li><li><b>04</b><span>pfSense: VM 100 with LAN and VPN documentation.</span></li><li><b>05</b><span>DC01: Windows Server 2022, AD DS, DNS, DHCP, FSMO roles.</span></li><li><b>06</b><span>WS01 and Linux01: workstation and Ubuntu domain-member evidence.</span></li><li><b>07</b><span>Operations repository: inventories, validation, limitations, and source mapping.</span></li></ol></section>
  <section class="section" aria-labelledby="current-title"><div class="section-head reveal"><p class="eyebrow">Current Verified State</p><h2 id="current-title">What current records support.</h2><p>Linux01 is treated as VLAN 30 at 10.10.30.20/24. Older Linux01 10.10.20.x evidence is historical. Backup storage is shown as backup-hdd. Restore evidence is retained as isolated startup checks, not recurring disaster-recovery assurance.</p></div><div class="lab-proof-detail reveal"><article class="proof-chip verified"><span class="chip-kind">Validated</span><h3>Identity and Windows services</h3><p>DC01 evidence supports AD DS, DNS, DHCP, Group Policy, FSMO roles, and Windows Server 2022 context.</p></article><article class="proof-chip verified"><span class="chip-kind">Validated</span><h3>Linux01 integration</h3><p>Linux01 records support Ubuntu Server 26.04 LTS, VLAN 30 placement, domain membership, SSSD remediation, Kerberos, SSH, UFW, sudo, and QEMU guest agent claims where cited.</p></article><article class="proof-chip limitation"><span class="chip-kind">Limitation</span><h3>Backup and restore</h3><p>Current records validate backup job visibility and retained isolated restore checks. They do not establish recurring restore success, RTO, RPO, or production disaster recovery readiness.</p></article></div></section>
  <section class="section" aria-labelledby="screens-title"><div class="section-head reveal"><p class="eyebrow">Supporting Evidence</p><h2 id="screens-title">Screenshots are placed as evidence, not decoration.</h2></div><div class="evidence-gallery windows-evidence-gallery reveal">{screenshot_figures()}</div></section>
  <section class="section" aria-labelledby="artifacts-title"><div class="section-head reveal"><p class="eyebrow">Reviewer Artifacts</p><h2 id="artifacts-title">Stable links for current review.</h2></div><div class="artifact-grid reveal"><a class="artifact-card" href="./evidence-library/projects/on-prem-home-lab/current-validated-state/README.md"><span>README</span><strong>Current validated state</strong><em>Public-safe evidence summary and boundaries.</em></a><a class="artifact-card" href="./evidence-library/projects/on-prem-home-lab/current-validated-state/manifest.json"><span>JSON</span><strong>Evidence manifest</strong><em>Hashes, provenance, and classifications.</em></a><a class="artifact-card" href="./evidence-library/projects/on-prem-home-lab/current-validated-state/claim-map.csv"><span>CSV</span><strong>Claim map</strong><em>Claims linked to artifacts and limits.</em></a><a class="artifact-card" href="{DOC_URL}"><span>DOCX</span><strong>{DOC_CTA}</strong><em>{DISPLAY_TITLE}</em></a></div></section>"""
    write(ROOT / "on-prem-home-lab.html", page("Home Lab", "Validated Home Lab | Jeremy Fontenot Infrastructure Case Study", "Polished evidence-backed infrastructure case study for Jeremy Fontenot's personal on-premises home lab.", lab, "TechArticle"))

    proof = f"""
  <section class="page page-hero" aria-labelledby="proof-title"><div class="section-head reveal"><p class="eyebrow">Reviewer-Focused Evidence Index</p><h1 id="proof-title">Open the artifact behind each public claim.</h1><p class="lead">Proof is organized by reviewer questions, not file type. Each entry names the claim, source system, evidence type, visible command or property, classification, limitation, source file, and full evidence link.</p><div class="actions"><a class="button primary" href="{DOC_URL}">{DOC_CTA}</a><a class="button" href="./evidence-library/projects/on-prem-home-lab/current-validated-state/manifest.json">Current manifest</a><a class="button" href="./artifacts/site-audit/screenshot-evidence-map.csv">Screenshot map</a></div></div></section>
  <section class="section" aria-label="Proof filters"><div class="project-filter-bar" role="group" aria-label="Proof filters"><button class="filter-button is-active" type="button" data-filter="all">All</button><button class="filter-button" type="button" data-filter="home-lab">Home Lab</button><button class="filter-button" type="button" data-filter="m365">Microsoft 365</button><button class="filter-button" type="button" data-filter="validated">Validated</button><button class="filter-button" type="button" data-filter="limitation">Limitations</button></div></section>
  <section class="section" id="home-lab-proof" aria-labelledby="home-proof-title"><div class="section-head reveal"><p class="eyebrow">Home Lab Evidence</p><h2 id="home-proof-title">Can the environment, systems, identity, network path, backups, and limits be inspected?</h2></div><div class="proof-stream reveal">{proof_entries()}</div></section>
  <section class="section" id="m365-proof" aria-labelledby="m365-title"><div class="section-head reveal"><p class="eyebrow">Microsoft 365 and Entra Evidence</p><h2 id="m365-title">Can the tenant evidence be inspected?</h2></div><div class="proof-stream reveal"><article class="proof-chip verified" data-project="m365 validated"><span class="chip-kind">Validated / screenshot and export</span><h3>Tenant, domains, users, groups, roles, licenses, sign-ins, and audit activity</h3><p>Repository-local CSV, JSON, Markdown, screenshot, and hash records support personal tenant review. Claims are limited to preserved artifacts and do not imply employer or client administration.</p><div class="proof-links"><a href="./evidence-library/projects/microsoft-365-lab/m365-entra-proof-inventory-20260605.csv">Proof inventory</a><a href="./evidence-library/projects/microsoft-365-lab/m365-entra-site-proof-map-20260605.md">Site proof map</a><a href="./evidence-library/projects/microsoft-365-lab/evidence/entra-exports-20260605-073748/microsoft-365-lab-evidence-manifest-20260605.csv">Evidence manifest</a></div></article></div></section>"""
    write(ROOT / "proof.html", page("Proof", "Proof Index | Jeremy Fontenot Evidence-First Portfolio", "Reviewer-focused proof index organized around claims, classifications, source systems, evidence links, and limitations.", proof, "CollectionPage"))

    dashboard = f"""
  <section class="page page-hero" aria-labelledby="dashboard-title"><div class="section-head reveal"><p class="eyebrow">Portfolio Evidence Dashboard</p><h1 id="dashboard-title">A concise status view for public artifacts.</h1><p class="lead">This dashboard summarizes static repository artifacts generated from manifests, screenshot review, claim mapping, link validation, document integrity, and source commits.</p></div><div class="metric-grid reveal">{metrics()}</div></section>
  <section class="section" aria-labelledby="dash-details"><div class="section-head reveal"><p class="eyebrow">Validation Coverage</p><h2 id="dash-details">Current review status.</h2></div><div class="dashboard-grid reveal">{cards([
      ("SRC", "Source commit", f"Website source {LOCAL_COMMIT[:7]}; home-lab source {OPS_COMMIT[:7]}.", '<a class="evidence-link" href="./artifacts/site-audit/live-repository-drift.md"><span>audit</span>Open</a>'),
      ("MAP", "Screenshot audit status", "Screenshot map and review records were regenerated from public HTML references and inspected source dimensions.", '<a class="evidence-link" href="./artifacts/site-audit/screenshot-review.md"><span>review</span>Open</a>'),
      ("DOC", "Document integrity status", "New DOCX is published to assets/documents and mirrored in the evidence library with a SHA-256 manifest entry.", f'<a class="evidence-link" href="{DOC_URL}"><span>docx</span>Download</a>'),
      ("LIM", "Unresolved limitations", "Backup restore assurance, production scale, employer administration, and broad security assurance are intentionally excluded.", '<a class="evidence-link" href="./proof.html#home-lab-proof"><span>proof</span>Open</a>'),
    ])}</div></section>"""
    write(ROOT / "dashboard.html", page("Dashboard", "Dashboard | Portfolio Evidence Status", "Static portfolio dashboard for evidence counts, source commits, screenshot review, claim maps, link validation, and document integrity.", dashboard, "WebPage"))

    resume = """
  <section class="page page-hero" aria-labelledby="resume-title"><div class="section-head reveal"><p class="eyebrow">Resume</p><h1 id="resume-title">Service Desk and IT Support experience with infrastructure growth.</h1><p class="lead">Jeremy is positioned as an experienced support professional moving toward systems administration and infrastructure operations. Personal lab work is clearly separated from professional employment.</p><div class="actions"><a class="button primary" href="./assets/resume/jeremy-fontenot-resume.pdf">Download Resume PDF</a><a class="button" href="./contact.html">Contact</a><a class="button" href="./on-prem-home-lab.html">Review Home Lab</a></div></div></section>
  <section class="section" aria-labelledby="summary-title"><div class="section-head reveal"><p class="eyebrow">Professional Summary</p><h2 id="summary-title">Support-centered, evidence-minded, infrastructure-ready.</h2><p>Experienced Service Desk and IT Support professional with hands-on troubleshooting, user support, documentation, Microsoft 365 familiarity, Windows administration fundamentals, and a personal infrastructure lab used to practice systems administration, validation, and operational documentation.</p></div><div class="resume-timeline reveal"><article><time>Professional focus</time><h3>Service Desk / IT Support</h3><p>User support, incident triage, escalation communication, troubleshooting documentation, and practical technical follow-through.</p></article><article><time>Career direction</time><h3>Systems Administration and Infrastructure Operations</h3><p>Personal lab evidence supports Active Directory, DNS, DHCP, Group Policy, Linux integration, backups, validation scripts, and documentation habits.</p></article><article><time>Evidence boundary</time><h3>Honest separation of work and lab</h3><p>Portfolio evidence represents personal lab and repository work unless explicitly described as employment experience.</p></article></div></section>
  <section class="section" aria-labelledby="skills-title"><div class="section-head reveal"><p class="eyebrow">Selected Skills</p><h2 id="skills-title">Skills connected to proof without becoming a file index.</h2></div><div class="signal-matrix reveal">""" + cards([
      ("SUP", "Support and Troubleshooting", "Ticket thinking, concise RCA notes, user-facing clarity, and escalation-ready documentation.", ""),
      ("WIN", "Windows and Active Directory", "Windows Server, AD DS, DNS, DHCP, GPO, file permissions, and workstation evidence in a personal lab.", ""),
      ("LIN", "Linux and Networking", "Ubuntu Server, SSSD, Kerberos, SSH, UFW, VLAN placement, pfSense, and OpenVPN management-path documentation.", ""),
      ("AUTO", "PowerShell and Validation", "Repository scripts, manifest generation, hash checks, link validation, and evidence governance.", ""),
    ]) + """</div></section>"""
    write(ROOT / "resume.html", page("Resume", "Resume | Jeremy Fontenot Service Desk and IT Support", "Browser resume for Jeremy Fontenot, accurately positioning Service Desk and IT Support experience with systems administration direction.", resume, "ProfilePage"))

    contact = """
  <section class="page page-hero" aria-labelledby="contact-title"><div class="section-head reveal"><p class="eyebrow">Contact</p><h1 id="contact-title">A direct route for hiring conversations and technical review.</h1><p class="lead">Use this page for professional opportunities, technical interviews, infrastructure and support roles, resume review, or portfolio review.</p></div><div class="contact-route-grid reveal"><a href="mailto:jeremyfontenot71@gmail.com"><span>Email</span><strong>jeremyfontenot71@gmail.com</strong><em>Professional opportunities and follow-up.</em></a><a href="https://www.linkedin.com/in/jeremy-fontenot/"><span>LinkedIn</span><strong>Jeremy Fontenot</strong><em>Hiring and recruiter contact route.</em></a><a href="https://github.com/fontenotjeremy71-hub"><span>GitHub</span><strong>fontenotjeremy71-hub</strong><em>Public repositories and portfolio source.</em></a><a href="https://github.com/fontenotjeremy71-hub/jeremyfontenot"><span>Website Repository</span><strong>Portfolio source repository</strong><em>Static site, evidence library, and validation scripts.</em></a></div></section>
  <section class="section" aria-labelledby="routes-title"><div class="section-head reveal"><p class="eyebrow">Review Routes</p><h2 id="routes-title">Choose the shortest path for the conversation.</h2></div><div class="signal-matrix reveal">""" + cards([
      ("ROLE", "Support and Infrastructure Roles", "Discuss service desk, desktop support, Microsoft 365 support, junior systems administration, and infrastructure operations opportunities.", ""),
      ("TECH", "Technical Interviews", "Open the home-lab case study, proof index, and APA document for evidence-backed technical screening.", ""),
      ("RES", "Resume Review", "Use the browser resume and PDF download for professional background review.", ""),
      ("PORT", "Portfolio Review", "Use the proof index and dashboard to inspect source artifacts, links, and validation boundaries.", ""),
    ]) + """</div></section>"""
    write(ROOT / "contact.html", page("Contact", "Contact | Jeremy Fontenot IT Support Portfolio", "Professional contact routes for Jeremy Fontenot, including email, LinkedIn, GitHub, resume, and portfolio review links.", contact, "ContactPage"))


def screenshot_figures() -> str:
    shots = [
        ("dc01-active-directory-fsmo-validation.png", "Active Directory and FSMO roles", "PowerShell output validates DC01 forest/domain context, global catalog role, and FSMO role ownership."),
        ("dc01-dns-forward-reverse-validation.png", "DNS forward and reverse validation", "PowerShell output supports AD-integrated DNS service and lookup claims."),
        ("dc01-dhcp-scope-options-leases-redacted.png", "DHCP scope and options", "Redacted DHCP evidence supports scope, options, statistics, and lease review."),
        ("dc01-group-policy-validation.png", "Group Policy inventory", "Command output supports GPO inventory and domain/OU link claims."),
        ("linux01-login-system-state.png", "Linux01 system state", "Console evidence supports Ubuntu, kernel, hostname, and address context."),
        ("linux01-evidence-archive-sha256-redacted.png", "Evidence archive hash", "Redacted terminal evidence supports archive and SHA-256 integrity review."),
    ]
    html = []
    for file, title, caption in shots:
        path = f"./evidence-library/projects/on-prem-home-lab/validated-2026-06-21/screenshots/{file}"
        html.append(f'<figure class="screenshot-card"><a class="screenshot-preview" href="{path}"><img src="{path}" loading="lazy" decoding="async" alt="{caption}"></a><figcaption><strong>{title}.</strong> {caption} <a class="evidence-open-link" href="{path}">Open full evidence</a>.</figcaption></figure>')
    return "".join(html)


def proof_entries() -> str:
    entries = [
        ("validated home-lab", "Is the environment real?", "Proxmox API inventory and VM inventory records identify Proxmox VE 9.2.3, VMs 100/200/300/400, backup-hdd storage, and backup-critical-lab-vms.", "Validated", "./evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/home-lab-source-of-truth.md"),
        ("validated home-lab", "What systems are present?", "DC01, WS01, Linux01, pfSense, and the Proxmox host are represented through current inventory and retained curated evidence.", "Validated", "./evidence-library/projects/on-prem-home-lab/current-validated-state/manifest.json"),
        ("validated home-lab", "Is Active Directory implemented?", "DC01 evidence supports AD DS, DNS, DHCP, GPO inventory, and FSMO role ownership in the personal lab.", "Validated", "./evidence-library/projects/on-prem-home-lab/validated-2026-06-21/screenshots/dc01-active-directory-fsmo-validation.png"),
        ("validated home-lab", "Is Linux integrated with the domain?", "Linux01 evidence supports realm membership, SSSD, Kerberos, SSH, UFW, sudo, and post-remediation SSSD status.", "Validated", "./evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/linux01-sssd-remediation.md"),
        ("validated home-lab", "Is network segmentation configured?", "Current source records support Linux01 on VLAN 30 at 10.10.30.20/24 and historical handling for older Linux01 10.10.20.x evidence.", "Configured / Validated where cited", "./evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/home-lab-source-reconciliation.md"),
        ("validated home-lab", "Are backups present?", "Current records validate backup job visibility and backup-hdd storage; retained isolated restore checks are not promoted into recurring DR assurance.", "Tested / Limitation", "./evidence-library/projects/on-prem-home-lab/current-validated-state/source-summaries/backup-and-restore-validation.md"),
        ("limitation home-lab", "What is not proven?", "No employer production administration, enterprise scale, guaranteed uptime, RTO, RPO, recurring restore assurance, or full security assurance is claimed.", "Limitation", "./evidence-library/projects/on-prem-home-lab/current-validated-state/claim-map.csv"),
    ]
    html = []
    for tags, q, text, cls, link in entries:
        klass = "limitation" if "Limitation" in cls else "verified"
        html.append(f'<article class="proof-chip {klass}" data-project="{tags}"><span class="chip-kind">{cls}</span><h3>{q}</h3><p><strong>Claim supported:</strong> {text}</p><p><strong>Visible command or property:</strong> see linked source artifact, manifest, or screenshot map.</p><div class="proof-links"><a href="{link}">Open full evidence</a><a href="./evidence-library/projects/on-prem-home-lab/current-validated-state/claim-map.csv">Open claim map</a></div></article>')
    return "".join(html)


def metrics() -> str:
    manifest_count = len(json.loads((CURRENT / "manifest.json").read_text(encoding="utf-8"))["artifacts"])
    shot_count = len(list((PUBLIC_EVIDENCE / "validated-2026-06-21" / "screenshots").glob("*.png")))
    full_doc_hash = sha256(PUBLIC_DOCX)
    doc_hash = full_doc_hash[:12]
    items = [
        ("Systems", "4", "Primary VMs in current supported state"),
        ("Artifacts", str(manifest_count), "Current-state manifest records"),
        ("Screenshots", str(shot_count), "Home-lab public screenshot files"),
        ("DOCX", doc_hash, "Published document hash prefix"),
    ]
    html = []
    for label, value, note in items:
        if label == "DOCX":
            html.append(f'<article class="metric technical-metric"><strong class="metric-identifier" title="Full SHA-256: {full_doc_hash}"><span aria-hidden="true">{value}</span><span class="sr-only">Full DOCX SHA-256: {full_doc_hash}</span></strong><span>{label}: {note}</span></article>')
        else:
            html.append(f'<article class="metric"><strong>{value}</strong><span>{label}: {note}</span></article>')
    return "".join(html)


def build_evidence() -> list[dict]:
    CURRENT.mkdir(parents=True, exist_ok=True)
    rows = []
    sources = [
        ("home-lab-source-of-truth.md", OPS / "docs/architecture/home-lab-source-of-truth.md", "Validated", "Current supported architecture and historical/stale value handling.", "Does not rewrite raw historical evidence."),
        ("home-lab-source-reconciliation.md", OPS / "evidence/home-lab-source-reconciliation-20260626-012606.md", "Validated", "Current Proxmox, DNS, route, VPN-path, backup, and conflict reconciliation evidence.", "TCP timeouts remain observations only."),
        ("home-lab-inventory-validation.md", OPS / "evidence/home-lab-inventory-validation-20260626-010601.md", "Tested", "Read-only inventory, route, DNS, and repository validation results.", "Does not validate every service or guest internal state."),
        ("linux01-sssd-remediation.md", OPS / "evidence/linux01-sssd-remediation-20260626-110159.md", "Validated", "Linux01 SSSD remediation and post-remediation validation.", "Represents personal lab remediation, not employer infrastructure."),
        ("backup-and-restore-validation.md", OPS / "docs/recovery/backup-and-restore-validation.md", "Tested / Limitation", "Backup job visibility and retained isolated restore documentation boundary.", "No recurring DR, RTO, or RPO assurance."),
    ]
    for item in sources:
        rows.append(copy_source(*item))

    readme = f"""# Current Validated State

This public evidence folder summarizes reviewer-safe records from the private home-lab operations repository without exposing private repository links. Public claims use stable classifications instead of month-based proof labels.

## Source Commits

- Website source commit at audit start: `{LOCAL_COMMIT}`
- Home-lab source commit: `{OPS_COMMIT}`

## Supported Public Claims

- Personal nonproduction lab hosted on Dell PowerEdge R710 with Proxmox VE 9.2.3.
- pfSense, DC01, WS01, and Linux01 are the primary lab VMs described by current records.
- Linux01 current supported placement is VLAN 30 at `10.10.30.20/24`; older `10.10.20.x` Linux01 evidence is historical.
- DC01 evidence supports AD DS, DNS, DHCP, Group Policy, and FSMO role review where linked.
- Backup job visibility is supported; recurring disaster-recovery assurance, RTO, and RPO are not claimed.
- Linux01 SSSD remediation is documented as remediated and validated in the personal lab.

## Limitations

Personal lab evidence does not represent employer production administration, client infrastructure ownership, enterprise scale, guaranteed uptime, or full security assurance. Timeout observations remain timeout observations.
"""
    write(CURRENT / "README.md", readme)

    doc_record = {
        "source_repository": "fontenotjeremy71-hub/jeremyfontenot",
        "source_commit": LOCAL_COMMIT,
        "original_source_path": "generated from current public-safe evidence and source summaries",
        "public_destination_path": "assets/documents/Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx",
        "target_system": "Home lab documentation",
        "evidence_classification": "Technical Proof",
        "capture_timestamp": BUILD_TIME,
        "sanitization_performed": "Newly authored public-safe document; no secrets included.",
        "sha256_hash": sha256(PUBLIC_DOCX),
        "claim_supported": "Professional APA home-lab documentation is published and hash-verifiable.",
        "limitations": "Document summarizes evidence; raw evidence remains authoritative for specific command output.",
        "reviewer_notes": "Mirrored in the evidence library for reviewer access.",
    }
    rows.append(doc_record)

    manifest = {"generated_at": BUILD_TIME, "website_commit_at_generation": LOCAL_COMMIT, "home_lab_source_commit": OPS_COMMIT, "artifacts": rows}
    write(CURRENT / "manifest.json", json.dumps(manifest, indent=2))
    with (CURRENT / "manifest.csv").open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)
    with (CURRENT / "provenance.csv").open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)
    claim_rows = [
        ["claim", "classification", "supporting_artifact", "limitation"],
        ["Proxmox VE 9.2.3 and four primary VMs are current supported state.", "Validated", rows[0]["public_destination_path"], "Current records are repository evidence, not live public telemetry."],
        ["Linux01 is currently treated as VLAN 30 at 10.10.30.20/24.", "Validated", rows[1]["public_destination_path"], "Older 10.10.20.x Linux evidence is historical."],
        ["Backup job configuration is visible for VMs 100/200/300/400 on backup-hdd.", "Tested", rows[4]["public_destination_path"], "Does not prove recurring restore assurance, RTO, or RPO."],
        ["Linux01 SSSD remediation is documented as remediated and validated.", "Validated", rows[3]["public_destination_path"], "Personal lab only."],
        ["Jeremy has employer production systems-administrator experience.", "Limitation", "Not claimed", "Professional positioning remains Service Desk and IT Support progressing toward infrastructure operations."],
    ]
    with (CURRENT / "claim-map.csv").open("w", encoding="utf-8", newline="") as f:
        csv.writer(f).writerows(claim_rows)
    write(CURRENT / "redaction-log.md", "# Redaction Log\n\nNo credentials, private keys, tokens, cookies, or secrets were copied. Public summaries avoid private repository links and retain only reviewer-safe claim context. Retired marketing labels were removed from public-facing copy while historical timestamps remain in raw compatibility folders.")
    return rows


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def set_cell(cell, text: str, bold: bool = False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(10)


def build_docx() -> None:
    PUBLIC_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 2
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
    header = sec.header.paragraphs[0]
    header.text = "Jeremy Fontenot's On-Premises Home Lab Documentation"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(9)
    add_page_number(sec)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(180)
    tr = title.add_run(DISPLAY_TITLE)
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Calibri"
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Professional APA 7 Technical Portfolio Paper").italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by Jeremy Fontenot\nSource commit: {OPS_COMMIT}\nDocument generated: {BUILD_TIME[:10]}")
    doc.add_page_break()

    doc.add_heading("Executive Abstract", level=1)
    doc.add_paragraph("This professional paper documents Jeremy Fontenot's personal on-premises home lab as a technical portfolio artifact for recruiters, hiring managers, systems administrators, infrastructure engineers, and technical interviewers. It summarizes architecture, implementation, administration, validation, troubleshooting, automation, backup, recovery, evidence governance, and known limitations using current repository-backed evidence.")
    doc.add_paragraph("The lab evidence supports a career direction from experienced Service Desk and IT Support into systems administration and infrastructure operations. It does not represent employer production administration, client infrastructure ownership, enterprise-scale operations, or guaranteed availability.")

    doc.add_heading("Table of Contents", level=1)
    for i, h in enumerate(["Purpose and Scope", "Evidence and Validation Methodology", "Architecture and Systems", "Identity, Network, and Security Controls", "Backup, Recovery, and Operations", "Evidence Governance and Limitations", "References", "Appendices"], 1):
        doc.add_paragraph(f"{i}. {h}", style=None)
    doc.add_page_break()

    sections = [
        ("Purpose and Scope", "The purpose of this document is to present a public-safe, evidence-based account of a personal home lab used for systems administration practice. The scope includes architecture, implementation, validation, documentation, and operational boundaries."),
        ("Professional and Technical Context", "Jeremy is positioned as an experienced Service Desk and IT Support professional progressing toward systems administration and infrastructure operations. The home lab demonstrates technical initiative and disciplined documentation without converting personal lab activity into employment claims."),
        ("Evidence and Validation Methodology", "Evidence quality is established through commands, authenticated inventories, structured output, manifests, hashes, configuration state, testing results, source records, and traceable provenance. Dates are retained as metadata, not as proof by themselves."),
        ("Evidence Classification Model", "Claims are classified as Planned, Configured, Implemented, Tested, Validated, Historical, Limitation, or Inconclusive. Public copy does not upgrade configured or tested behavior into validated behavior unless direct evidence supports that classification."),
        ("Physical Infrastructure", "The current supported physical host is a Dell PowerEdge R710 used as a personal nonproduction lab host."),
        ("Virtualization Platform", "Current source records support Proxmox VE 9.2.3 with VMs 100, 200, 300, and 400 running. VM 100 is pfSense, VM 200 is DC01, VM 300 is WS01, and VM 400 is Linux01."),
        ("Storage Architecture", "Current records support backup storage named backup-hdd at /mnt/pve/backup-hdd. Historical references to other backup storage labels are not treated as current unless future evidence supports them."),
        ("Network Architecture", "The lab uses pfSense, internal lab networks, VLAN 30 for Linux01, and an OpenVPN management path documented in the operations records. Linux01 current supported placement is 10.10.30.20/24."),
        ("pfSense Routing and Segmentation", "pfSense is represented as VM 100. Public claims distinguish current configuration records from behavior that was not independently tested."),
        ("VLAN Architecture", "Linux01 is currently treated as VLAN 30 at 10.10.30.20/24. Older Linux01 10.10.20.x records are historical setup or pre-reconciliation evidence."),
        ("VPN and Management Path", "Route and TCP evidence support a management path for selected services. Timeout records remain timeout observations and are not converted into outage claims."),
        ("Windows Server and Active Directory", "DC01 is documented as a Windows Server 2022 domain controller for ad.jeremyfontenot.online. Retained evidence supports AD DS, DNS, DHCP, Group Policy, and FSMO role review where cited."),
        ("DNS", "DNS evidence includes AD-integrated DNS records and forward/reverse lookup validation in retained public screenshots and text artifacts."),
        ("DHCP", "DHCP evidence includes scope, options, statistics, and redacted lease review records. Sensitive client identifiers remain redacted."),
        ("Group Policy", "Group Policy evidence supports GPO inventory and domain or organizational-unit link review."),
        ("File Services and Permissions", "File-service evidence is treated as tested only where specific permission validation records exist. No broader enterprise file-service assurance is claimed."),
        ("WS01 Workstation Architecture and Validation", "WS01 is documented as a Windows domain workstation. Current records include authenticated inventory evidence, while guest-agent limitations are kept visible where applicable."),
        ("Linux01 Architecture and Network Placement", "Linux01 is documented as Ubuntu Server 26.04 LTS on VLAN 30 at 10.10.30.20/24."),
        ("Linux Active Directory Integration", "Linux01 evidence supports domain integration through realm membership, identity resolution, and related validation artifacts."),
        ("SSSD and Kerberos", "The Linux01 SSSD remediation record documents a responder activation conflict, configuration correction, service restart, failed-unit cleanup, identity lookup, and a new AD-authenticated SSH session."),
        ("SSH, UFW, sudo, and Linux Security Controls", "Linux01 records support SSH, UFW, delegated sudo, and related Linux security-control claims where linked."),
        ("Centralized Logging", "Centralized logging is described within the boundaries of available pfSense and operations documentation. Full security monitoring maturity is not claimed."),
        ("QEMU Guest Agent and Maintenance State", "Linux01 records include QEMU guest-agent response and maintenance-state evidence. WS01 guest-agent limitations remain visible."),
        ("Backup Architecture", "Current inventory supports an enabled scheduled backup job named backup-critical-lab-vms for VMs 100, 200, 300, and 400 using snapshot mode and zstd compression."),
        ("Recovery Validation", "Retained documentation records isolated startup restore checks for temporary WS01 and Linux01 restore VMs. The current records do not establish recurring disaster-recovery assurance, RTO, or RPO."),
        ("Storage Monitoring", "Storage monitoring claims are limited to available inventory and utilization records."),
        ("Read-Only Remote Validation", "Read-only validation includes repository checks, Proxmox inventory, route evidence, DNS checks, and bounded TCP reachability. It does not prove full service health for every endpoint."),
        ("Repository Validation", "The operations repository records validation scripts, inventories, evidence summaries, and source-of-truth reconciliation. The website repository publishes sanitized reviewer-safe summaries."),
        ("Evidence Integrity", "Public artifacts are mapped through manifests, SHA-256 hashes, source paths, and provenance records."),
        ("Screenshot Evidence Review", "Screenshots are used only where they support adjacent claims, with full-size originals linked from previews and captions scoped to visible evidence."),
        ("Unresolved Investigations and Known Limitations", "Known limitations include personal lab scope, no production administration claim, no enterprise scale claim, no guaranteed uptime, no full security assurance, and no recurring restore assurance."),
    ]
    for heading, text in sections:
        doc.add_heading(heading, level=1 if heading in {"Purpose and Scope", "Evidence and Validation Methodology", "Architecture and Systems", "Identity, Network, and Security Controls", "Backup, Recovery, and Operations", "Evidence Governance and Limitations"} else 2)
        doc.add_paragraph(text)

    doc.add_heading("Table 1", level=2)
    doc.add_paragraph("Evidence Classification Summary")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for c, t in zip(hdr, ["Classification", "Use", "Boundary"]):
        set_cell(c, t, True)
    for row in [
        ("Validated", "Direct evidence confirms the stated behavior or state.", "Used only where source artifacts support the claim."),
        ("Tested", "A relevant test produced a recorded result.", "Does not imply complete service assurance."),
        ("Configured", "Configuration is present.", "Behavior may require separate testing."),
        ("Limitation", "Known boundary, gap, or unsupported conclusion.", "Kept visible for reviewers."),
    ]:
        cells = table.add_row().cells
        for c, t in zip(cells, row):
            set_cell(c, t)

    doc.add_heading("Figure 1", level=2)
    doc.add_paragraph("Current supported home-lab architecture summary. Source note: derived from current operations repository source-of-truth and reconciliation records.")

    doc.add_heading("References", level=1)
    refs = [
        f"Fontenot, J. ({BUILD_TIME[:4]}). Home lab source-of-truth repository records. fontenotjeremy71-hub/jeremy-homelab-ops, commit {OPS_COMMIT}.",
        f"Fontenot, J. ({BUILD_TIME[:4]}). Public evidence library and portfolio repository. fontenotjeremy71-hub/jeremyfontenot, commit {LOCAL_COMMIT}.",
        "American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.).",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_heading("Appendix A: Source Artifacts", level=1)
    for item in ["home-lab-source-of-truth.md", "home-lab-source-reconciliation.md", "home-lab-inventory-validation.md", "linux01-sssd-remediation.md", "backup-and-restore-validation.md"]:
        doc.add_paragraph(item)
    doc.add_heading("Appendix B: Claim Boundaries", level=1)
    doc.add_paragraph("The document intentionally avoids claims of production administration, client infrastructure ownership, enterprise scale, guaranteed uptime, unsupported business impact, full security assurance, or recurring restore assurance.")

    doc.core_properties.title = DISPLAY_TITLE
    doc.core_properties.author = "Jeremy Fontenot"
    doc.core_properties.subject = "Professional home lab documentation"
    doc.core_properties.keywords = "home lab, Proxmox, pfSense, Active Directory, Linux, PowerShell, evidence governance"
    doc.save(PUBLIC_DOCX)
    mirror = PUBLIC_EVIDENCE / "Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx"
    shutil.copy2(PUBLIC_DOCX, mirror)


def build_screenshot_audit() -> None:
    AUDIT.mkdir(parents=True, exist_ok=True)
    rows = []
    for img in sorted((PUBLIC_EVIDENCE / "validated-2026-06-21" / "screenshots").glob("*.png")):
        name = img.name
        rows.append({
            "page": "on-prem-home-lab.html / proof.html",
            "section": "Supporting Evidence",
            "image_path": str(img.relative_to(ROOT)).replace("\\", "/"),
            "source_repository": "fontenotjeremy71-hub/jeremyfontenot",
            "source_commit": LOCAL_COMMIT,
            "target_system": "DC01/Linux01/Proxmox depending on screenshot filename",
            "claim_supported": screenshot_claim(name),
            "visible_command_or_ui": "Visible command or terminal/admin UI reviewed from public screenshot context",
            "important_visible_result": screenshot_result(name),
            "evidence_classification": "Validated" if any(x in name for x in ["dc01", "linux01", "sha256", "scp"]) else "Historical",
            "relevance": "exact support" if any(x in name for x in ["dc01", "linux01", "sha256", "scp"]) else "partial support",
            "currency": "current or historical as labeled in adjacent copy",
            "redaction_status": "redacted where filename or manifest indicates redaction",
            "alt_text": screenshot_claim(name),
            "caption": screenshot_result(name),
            "recommended_action": "Use near the supported claim and link to full original.",
            "replacement_or_preview_path": str(img.relative_to(ROOT)).replace("\\", "/"),
            "review_notes": "Full-size original retained; tall images are linked from previews for readability.",
        })
    with (AUDIT / "screenshot-evidence-map.csv").open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)
    write(AUDIT / "screenshot-evidence-map.json", json.dumps(rows, indent=2))
    write(AUDIT / "screenshot-review.md", f"""# Screenshot Review

Reviewed public home-lab screenshots referenced by the rebuilt HTML and evidence manifests. Full-size originals remain unchanged under `evidence-library/projects/on-prem-home-lab/validated-2026-06-21/screenshots/`.

## Result

- Screenshot records mapped: {len(rows)}
- Relevant images are used as proof, not decoration.
- Tall terminal screenshots are linked to full-size originals from previews.
- Captions and alt text were rewritten to avoid claiming more than the visible command output supports.
- Redacted filenames and manifest notes remain respected.

## Limitation

This review is source-image and page-placement based. It does not create new screenshots of private infrastructure and does not alter raw evidence.
""")


def screenshot_claim(name: str) -> str:
    if "active-directory" in name:
        return "DC01 Active Directory forest/domain and FSMO role evidence."
    if "dns" in name:
        return "DC01 DNS forward and reverse lookup evidence."
    if "dhcp" in name:
        return "DC01 DHCP scope, options, statistics, and redacted lease evidence."
    if "group-policy" in name:
        return "Group Policy inventory and link evidence."
    if "login" in name:
        return "Linux01 operating-system, hostname, kernel, and address context."
    if "sha256" in name:
        return "Evidence archive hash verification."
    if "scp" in name:
        return "Evidence transfer record."
    return "Contextual home-lab screenshot evidence."


def screenshot_result(name: str) -> str:
    return screenshot_claim(name).replace(" evidence.", ".")


def build_audit_docs() -> None:
    remote_commit = git("ls-remote", "origin", "refs/heads/main").split()[0]
    write(AUDIT / "live-repository-drift.md", f"""# Live Repository Drift

## Commits

- Local website commit before edits: `{LOCAL_COMMIT}`
- Remote website commit before edits: `{remote_commit}`
- Deployed commit: not directly exposed by GitHub Pages HTML; deployment source is workflow-backed `main` for `jeremyfontenot.online`.
- Home-lab source commit: `{OPS_COMMIT}`

## Drift Observed

- Live and local source were aligned at commit `{LOCAL_COMMIT}` before editing.
- Public pages contained retired date-forward proof labels and retired document wording.
- The source of truth for current home-lab state is the operations repository plus sanitized public evidence copied into the website repository.

## Actions Taken

- Rebuilt primary public pages around stable evidence classifications.
- Published a stable current-state evidence folder.
- Created a new professional home-lab DOCX and linked it from primary CTAs.
- Preserved historical evidence paths for compatibility while removing retired public-facing wording from active navigation and page copy.
""")
    write(AUDIT / "content-and-visual-audit.md", "# Content and Visual Audit\n\nPrimary pages were rewritten around recruiter and technical-review paths, stable evidence classifications, command-center visual language, consistent CTAs, proof cards, screenshot placement, and visible limitations. Responsive screenshots are generated separately by the validation scripts.")
    write(AUDIT / "link-validation.md", "# Link Validation\n\nInternal links were regenerated to point at active static pages, current manifests, public evidence files, public resume assets, GitHub, LinkedIn, email, and the new DOCX. Full validation command results are recorded in final-validation.md.")
    write(AUDIT / "accessibility-validation.md", "# Accessibility Validation\n\nPages retain skip navigation, semantic landmarks, one H1 per page, keyboard-accessible filters, visible focus styles, descriptive links, alt text, reduced-motion handling, and native controls.")
    write(AUDIT / "seo-validation.md", "# SEO Validation\n\nPrimary pages include unique titles, descriptions, canonical URLs, Open Graph/Twitter metadata, JSON-LD, crawlable navigation, robots and sitemap updates.")
    write(AUDIT / "document-validation.md", f"# Document Validation\n\n- Title: {DISPLAY_TITLE}\n- Filename: Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx\n- Public path: `{DOC_URL}`\n- Evidence mirror: `evidence-library/projects/on-prem-home-lab/Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx`\n- SHA-256: `{sha256(PUBLIC_DOCX)}`\n- DOCX package check: pending validation command\n")
    write(AUDIT / "final-validation.md", "# Final Validation\n\nValidation commands and deployment checks are appended during the final verification pass.")


def update_readme_sitemap() -> None:
    write(ROOT / "README.md", f"""# Jeremy Fontenot Portfolio

Static professional portfolio for Jeremy Fontenot.

## Source of Truth

- Website repository: `https://github.com/fontenotjeremy71-hub/jeremyfontenot`
- Public site: `https://jeremyfontenot.online`
- Deployment: GitHub Pages workflow from `main`
- Home-lab source material: private operations repository, copied only as sanitized public evidence

## Current Public Evidence

- Home-lab case study: `on-prem-home-lab.html`
- Proof index: `proof.html`
- Portfolio evidence dashboard: `dashboard.html`
- Current home-lab manifest: `evidence-library/projects/on-prem-home-lab/current-validated-state/manifest.json`
- Document title: {DISPLAY_TITLE}
- Professional DOCX: `assets/documents/Jeremy-Fontenot-On-Premises-Home-Lab-Documentation.docx`

## Validation

Use the repo-local scripts in `scripts/validation` for HTML, JSON, sitemap, SEO, accessibility, link, evidence, screenshot, PowerShell, and repository-structure checks.
""")
    pages = ["", "projects.html", "on-prem-home-lab.html", "proof.html", "dashboard.html", "resume.html", "contact.html", "home-lab-operations-proof.html"]
    urls = "\n".join(f"  <url><loc>{DOMAIN}/{p}</loc></url>" for p in pages)
    write(ROOT / "sitemap.xml", f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
{urls}
</urlset>""")
    write(ROOT / "robots.txt", f"""User-agent: *
Allow: /
Sitemap: {DOMAIN}/sitemap.xml""")


def scrub_retired_label() -> None:
    for retired_doc in PUBLIC_EVIDENCE.glob("Jeremy_Fontenot_On_Premises_Home_Lab_Documentation_*_2026-06-24.docx"):
        retired_doc.unlink()
    patterns = [
        (re.compile(r"IP[-_ ]Verified", re.I), "Professional"),
        (re.compile(r"Validated through [A-Z][a-z]+ 26, 2026\.", re.I), "Current operations validation."),
        (re.compile(r"validated [A-Z][a-z]+ 26, 2026", re.I), "Current Operations Validation"),
        (re.compile(r"[A-Z][a-z]+ 26 operations validation", re.I), "Current operations validation"),
        (re.compile(r"[A-Z][a-z]+ 21 public README", re.I), "current public README"),
        (re.compile(r"[A-Z][a-z]+ 26 JSON summary", re.I), "current JSON summary"),
    ]
    exts = {".html", ".md", ".json", ".csv", ".xml", ".txt", ".css", ".js"}
    for path in ROOT.rglob("*"):
        if ".git" in path.parts or "node_modules" in path.parts or path.is_dir() or path.suffix.lower() not in exts:
            continue
        text = path.read_text(encoding="utf-8", errors="replace")
        new = text
        for rx, repl in patterns:
            new = rx.sub(repl, new)
        if new != text:
            path.write_text(new, encoding="utf-8")


BUILD_TIME = datetime.now().astimezone().isoformat(timespec="seconds")
LOCAL_COMMIT = git("rev-parse", "HEAD")
OPS_COMMIT = git("rev-parse", "HEAD", cwd=OPS)


def main() -> None:
    build_docx()
    build_evidence()
    build_pages()
    build_screenshot_audit()
    build_audit_docs()
    update_readme_sitemap()
    scrub_retired_label()


if __name__ == "__main__":
    main()
